"""
build_doc.py — Generic document compiler for Troy J. Hokanson
==============================================================
Builds resumes, cover letters, and stationary letters.
Header, footer, and layout are ALWAYS injected via docx_header.py.
NEVER hand-roll the header here.

Usage:
    from build_doc import build_document
    from specs.example_spec import RESUME_SPEC, COVER_SPEC

    build_document("resume",       RESUME_SPEC, "output/Hokanson_Resume_Role.docx")
    build_document("cover_letter", COVER_SPEC,  "output/Hokanson_Cover_Role.docx")
    build_document("stationary",   LETTER_SPEC, "output/Hokanson_Letter_Role.docx")

Spec schema is defined in each function below.
"""

from datetime import date
from pathlib import Path
from docx import Document
from docx_header import (
    new_document,
    build_navy_header,
    build_footer,
    add_section_heading,
    add_job_block,
    add_bullet,
    set_run,
    set_paragraph_format,
    BLACK,
)
from config import TROY_NAME, TROY_PHONE, TROY_EMAIL, TROY_LOCATION, TROY_LINKEDIN

# --------------------------------------------------------------------
# Locked education constants — verbatim from EDUCATION_CONSTANTS.md
# DO NOT paraphrase, reorder, or abbreviate these lines.
# --------------------------------------------------------------------
EDUCATION_LINES = [
    "Master of Arts, Police Leadership, Administration and Education",
    "University of St. Thomas, St. Paul, MN",
    "GPA: 3.94",
    "2005",
    "",
    "Bachelor of Arts, Criminal Justice, Magna Cum Laude",
    "St. Cloud State University, St. Cloud, MN",
    "GPA: 3.51",
    "1998",
    "",
    "Associate of Arts, Criminal Justice, Magna Cum Laude",
    "St. Cloud State University, St. Cloud, MN",
    "1996",
]


# ====================================================================
# BODY BUILDERS — one per document type
# ====================================================================

def _build_resume_body(doc: Document, spec: dict) -> None:
    """
    spec keys:
        summary   : str
        skills    : str
        jobs      : list of {title, employer, dates, bullets: [str]}
        extra_sections : list of {heading, paragraphs: [str]} (optional)
    Education is always appended verbatim from EDUCATION_LINES.
    """
    add_section_heading(doc, "Professional Summary")
    p = doc.add_paragraph(spec["summary"])
    set_paragraph_format(p, before=0, after=4, line=1.15)
    set_run(p.runs[0], size=10.5)

    add_section_heading(doc, "Core Skills")
    p = doc.add_paragraph(spec["skills"])
    set_paragraph_format(p, before=0, after=4, line=1.15)
    set_run(p.runs[0], size=10.5)

    add_section_heading(doc, "Experience")
    for job in spec.get("jobs", []):
        add_job_block(doc, job["title"], job["employer"], job["dates"])
        for bullet in job.get("bullets", []):
            add_bullet(doc, bullet)

    for section in spec.get("extra_sections", []):
        add_section_heading(doc, section["heading"])
        for para in section["paragraphs"]:
            p = doc.add_paragraph(para)
            set_paragraph_format(p, before=0, after=4, line=1.15)
            set_run(p.runs[0], size=10.5)

    add_section_heading(doc, "Education")
    for line in EDUCATION_LINES:
        p = doc.add_paragraph(line)
        set_paragraph_format(p, before=0, after=0, line=1.0)
        if p.runs:
            set_run(p.runs[0], size=10.5)


def _build_cover_letter_body(doc: Document, spec: dict) -> None:
    """
    spec keys:
        date_str          : str  e.g. "June 15, 2026"
        recipient_name    : str  e.g. "Hiring Manager"
        recipient_org     : str  e.g. "3M"
        recipient_address : list of str
        salutation        : str  e.g. "Hiring Manager"
        body_paragraphs   : list of str
    Always closes with "Respectfully," and TROY_NAME.
    """
    sender_lines = [
        TROY_NAME,
        TROY_LOCATION,
        TROY_PHONE,
        TROY_EMAIL,
        TROY_LINKEDIN,
    ]
    for line in sender_lines:
        if line:
            p = doc.add_paragraph(line)
            set_paragraph_format(p, before=0, after=0, line=1.0)
            set_run(p.runs[0], size=10.5)

    doc.add_paragraph("")
    p = doc.add_paragraph(spec.get("date_str", date.today().strftime("%B %d, %Y")))
    set_paragraph_format(p, before=0, after=0, line=1.0)
    doc.add_paragraph("")

    for line in [spec.get("recipient_name", ""), spec.get("recipient_org", "")] + spec.get("recipient_address", []):
        if line:
            p = doc.add_paragraph(line)
            set_paragraph_format(p, before=0, after=0, line=1.0)
            set_run(p.runs[0], size=10.5)

    doc.add_paragraph("")
    p = doc.add_paragraph(f"Dear {spec.get('salutation', 'Hiring Manager')},")
    set_paragraph_format(p, before=0, after=6, line=1.15)
    set_run(p.runs[0], size=10.5)

    for para in spec.get("body_paragraphs", []):
        p = doc.add_paragraph(para)
        set_paragraph_format(p, before=0, after=6, line=1.15)
        set_run(p.runs[0], size=10.5)

    doc.add_paragraph("")
    p = doc.add_paragraph("Respectfully,")
    set_paragraph_format(p, before=0, after=0, line=1.0)
    set_run(p.runs[0], size=10.5)
    p = doc.add_paragraph(TROY_NAME)
    set_paragraph_format(p, before=0, after=0, line=1.0)
    set_run(p.runs[0], size=10.5)


def _build_stationary_body(doc: Document, spec: dict) -> None:
    """
    spec keys:
        date_str          : str  (optional, defaults to today)
        recipient_name    : str  (optional)
        recipient_org     : str  (optional)
        recipient_address : list of str (optional)
        salutation        : str  (optional)
        body_paragraphs   : list of str
        closing           : str  (optional, defaults to "Respectfully,")
    """
    doc.add_paragraph("")
    p = doc.add_paragraph(spec.get("date_str", date.today().strftime("%B %d, %Y")))
    set_paragraph_format(p, before=0, after=0, line=1.0)
    doc.add_paragraph("")

    for line in [spec.get("recipient_name", ""), spec.get("recipient_org", "")] + spec.get("recipient_address", []):
        if line:
            p = doc.add_paragraph(line)
            set_paragraph_format(p, before=0, after=0, line=1.0)
            set_run(p.runs[0], size=10.5)

    if spec.get("salutation"):
        doc.add_paragraph("")
        p = doc.add_paragraph(f"Dear {spec['salutation']},")
        set_paragraph_format(p, before=0, after=6, line=1.15)
        set_run(p.runs[0], size=10.5)

    for para in spec.get("body_paragraphs", []):
        p = doc.add_paragraph(para)
        set_paragraph_format(p, before=0, after=6, line=1.15)
        set_run(p.runs[0], size=10.5)

    doc.add_paragraph("")
    closing = spec.get("closing", "Respectfully,")
    p = doc.add_paragraph(closing)
    set_paragraph_format(p, before=0, after=0, line=1.0)
    set_run(p.runs[0], size=10.5)
    p = doc.add_paragraph(TROY_NAME)
    set_paragraph_format(p, before=0, after=0, line=1.0)
    set_run(p.runs[0], size=10.5)


# ====================================================================
# MAIN PUBLIC FUNCTION
# ====================================================================

_BUILDERS = {
    "resume":       _build_resume_body,
    "cover_letter": _build_cover_letter_body,
    "stationary":   _build_stationary_body,
}


def build_document(doc_type: str, spec: dict, out_path: str | Path) -> None:
    """
    Build a compliant DOCX with locked navy/gold header and footer.

    Args:
        doc_type : "resume" | "cover_letter" | "stationary"
        spec     : content dict for the chosen doc type (see builders above)
        out_path : file path to save the DOCX
    """
    if doc_type not in _BUILDERS:
        raise ValueError(f"doc_type must be one of {list(_BUILDERS)}. Got: {doc_type!r}")

    Path(out_path).parent.mkdir(parents=True, exist_ok=True)
    doc = new_document()
    build_navy_header(doc)
    _BUILDERS[doc_type](doc, spec)
    build_footer(doc, show_page_numbers=False)
    doc.save(out_path)
    print(f"Built {doc_type}: {out_path}")
