#!/usr/bin/env python3

from __future__ import annotations

import hashlib
import json
import re
import sys
import zipfile
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

REPO_ROOT = Path(__file__).resolve().parents[2]
sys.path.insert(0, str(REPO_ROOT))

import config
import docx_header
from docx_header import add_paragraph_bottom_border, build_navy_header

APP_DIR = Path(__file__).resolve().parent
GENERAL_DIR = REPO_ROOT / "applications" / "2026-08-17_general_online_resume"
OUTPUT_DIR = APP_DIR / "output"
LOG_DIR = APP_DIR / "build_logs"

FONT = "EB Garamond"
BLACK = RGBColor(0x14, 0x14, 0x14)
GRAY = RGBColor(0x55, 0x55, 0x55)
STEEL = RGBColor(0x2D, 0x6A, 0x9F)


def set_font(run, size=10.0, *, bold=False, italic=False, color=BLACK):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        r_fonts.set(qn(f"w:{attr}"), FONT)


def set_para(p, *, before=0, after=2, line=1.02, keep_next=False, keep_together=False):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    pf.keep_with_next = keep_next
    pf.keep_together = keep_together
    pf.widow_control = True


def add_rich_text(p, text, size=10.0, color=BLACK):
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        is_bold = part.startswith("**") and part.endswith("**")
        clean = part[2:-2] if is_bold else part
        run = p.add_run(clean)
        set_font(run, size=size, bold=is_bold, color=color)


def configure_document(*, branded, compact=False, dense=False):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(0.54 if dense else (0.62 if compact else 0.68))
    section.right_margin = Inches(0.54 if dense else (0.62 if compact else 0.68))
    section.bottom_margin = Inches(0.42 if dense else (0.52 if compact else 0.62))
    section.top_margin = Inches(0.42 if dense else 0.55)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(9.0 if dense else (9.6 if compact else 10.1))
    normal.paragraph_format.space_after = Pt(2)
    normal.paragraph_format.line_spacing = 1.02

    if branded:
        # The portfolio is intentionally excluded while its public facts are stale.
        docx_header.CONTACT_PARTS = [
            (config.TROY_PHONE, f"tel:{config.TROY_PHONE.replace('.', '').replace('-', '').replace(' ', '')}"),
            (config.TROY_EMAIL, f"mailto:{config.TROY_EMAIL}"),
            (
                config.TROY_LINKEDIN,
                f"https://www.{config.TROY_LINKEDIN}"
                if config.TROY_LINKEDIN and not config.TROY_LINKEDIN.startswith("http")
                else config.TROY_LINKEDIN,
            ),
        ]
        build_navy_header(
            doc,
            body_top_margin_inches=1.32 if dense else 1.45,
            body_bottom_margin_inches=section.bottom_margin.inches,
            body_left_margin_inches=section.left_margin.inches,
            body_right_margin_inches=section.right_margin.inches,
        )
    return doc


def add_plain_header(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para(p, after=1, line=1.0)
    set_font(p.add_run("Troy Hokanson"), size=18, bold=True)
    contact = " | ".join(x for x in (config.TROY_PHONE, config.TROY_EMAIL, config.TROY_LINKEDIN) if x)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para(p, after=5, line=1.0)
    set_font(p.add_run(contact), size=9.5)


def add_section(doc, text, compact, dense=False):
    p = doc.add_paragraph()
    set_para(p, before=4 if dense else (5 if compact else 7), after=1.5 if dense else 2, line=1.0, keep_next=True)
    set_font(p.add_run(text.upper()), size=10.0 if dense else (10.5 if compact else 11.2), bold=True, color=STEEL)
    add_paragraph_bottom_border(p, color_hex="C9A84C", size=5)


def add_markdown(doc, source, *, branded, compact, dense):
    lines = source.read_text(encoding="utf-8").splitlines()
    for raw in lines:
        line = raw.strip()
        if not line or line.startswith(">"):
            continue
        if line == "<!-- page-break -->":
            doc.add_page_break()
            continue
        if line.startswith("# "):
            if not branded:
                add_plain_header(doc)
            continue
        if line.startswith("## "):
            add_section(doc, line[3:], compact, dense)
            continue
        if line.startswith("### "):
            if dense and not branded and line[4:] == "Detective / Electronic Crimes Unit (ECU)":
                doc.add_page_break()
            p = doc.add_paragraph()
            set_para(p, before=3 if dense else 4, after=0, line=1.0, keep_next=True, keep_together=True)
            add_rich_text(p, line[4:], size=9.5 if dense else (10.0 if compact else 10.5))
            for run in p.runs:
                run.bold = True
            continue
        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            set_para(p, after=0.8 if dense else (1.3 if compact else 2), line=1.0 if compact else 1.03, keep_together=True)
            p.paragraph_format.left_indent = Inches(0.18)
            p.paragraph_format.first_line_indent = Inches(-0.18)
            add_rich_text(p, line[2:], size=8.9 if dense else (9.45 if compact else 9.9))
            continue
        p = doc.add_paragraph()
        set_para(p, after=1.5 if dense else (2.5 if compact else 4), line=1.0 if dense else (1.02 if compact else 1.06), keep_together=True)
        add_rich_text(p, line, size=9.0 if dense else (9.55 if compact else 10.05))


def save_doc(source, destination, *, branded, compact, core_properties):
    dense = source.parent == GENERAL_DIR
    doc = configure_document(branded=branded, compact=compact, dense=dense)
    add_markdown(doc, source, branded=branded, compact=compact, dense=dense)
    props = doc.core_properties
    props.title = core_properties["title"]
    props.subject = core_properties["subject"]
    props.author = "Troy Hokanson"
    props.keywords = core_properties["keywords"]
    props.comments = "Generated from verified repository source on 2026-08-30. Drafting status."
    doc.save(destination)
    scrub_disallowed_font_refs(destination)


def scrub_disallowed_font_refs(path):
    """Replace non-visible Word defaults that trigger the strict font gate."""
    temp = path.with_suffix(".scrubbed.docx")
    with zipfile.ZipFile(path, "r") as source_zip, zipfile.ZipFile(
        temp, "w", compression=zipfile.ZIP_DEFLATED
    ) as target_zip:
        for item in source_zip.infolist():
            data = source_zip.read(item.filename)
            if item.filename.endswith(".xml"):
                data = data.replace(b"Calibri", b"EB Garamond").replace(b"Arial", b"EB Garamond")
            target_zip.writestr(item, data)
    temp.replace(path)


def sha256(path):
    return hashlib.sha256(path.read_bytes()).hexdigest()


def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    LOG_DIR.mkdir(parents=True, exist_ok=True)
    date = "2026-08-30"
    outputs = [
        (
            GENERAL_DIR / "resume.md",
            OUTPUT_DIR / f"Troy_Hokanson_General_Resume_ATS_{date}.docx",
            False,
            True,
            {
                "title": "Troy Hokanson — General Resume (ATS)",
                "subject": "Investigations, digital forensics, fraud, and public-safety technology",
                "keywords": "investigations, digital forensics, fraud, OSINT, public safety technology",
            },
        ),
        (
            GENERAL_DIR / "resume.md",
            OUTPUT_DIR / f"Troy_Hokanson_General_Resume_Branded_{date}.docx",
            True,
            True,
            {
                "title": "Troy Hokanson — General Resume",
                "subject": "Investigations, digital forensics, fraud, and public-safety technology",
                "keywords": "investigations, digital forensics, fraud, OSINT, public safety technology",
            },
        ),
        (
            APP_DIR / "resume.md",
            OUTPUT_DIR / f"Troy_Hokanson_OpenAI_Protective_Intelligence_Resume_{date}.docx",
            True,
            True,
            {
                "title": "Troy Hokanson — OpenAI Protective Intelligence & Threat Analyst Resume",
                "subject": "Application for Protective Intelligence & Threat Analyst",
                "keywords": "protective intelligence, threat analysis, OSINT, digital evidence, investigations",
            },
        ),
        (
            APP_DIR / "cover_letter.md",
            OUTPUT_DIR / f"Troy_Hokanson_OpenAI_Protective_Intelligence_Cover_Letter_{date}.docx",
            True,
            False,
            {
                "title": "Troy Hokanson — OpenAI Protective Intelligence & Threat Analyst Cover Letter",
                "subject": "Application for Protective Intelligence & Threat Analyst",
                "keywords": "OpenAI, protective intelligence, threat analysis, OSINT",
            },
        ),
    ]
    for source, destination, branded, compact, props in outputs:
        save_doc(source, destination, branded=branded, compact=compact, core_properties=props)

    manifest = {
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "status": "Drafting",
        "portfolio_omitted": True,
        "sources": {str(source.relative_to(REPO_ROOT)): sha256(source) for source, *_ in outputs},
        "outputs": {path.name: sha256(path) for path in OUTPUT_DIR.glob("*.docx")},
    }
    (LOG_DIR / "provenance.json").write_text(json.dumps(manifest, indent=2) + "\n", encoding="utf-8")
    print(json.dumps(manifest, indent=2))


if __name__ == "__main__":
    main()
