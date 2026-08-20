#!/usr/bin/env python3

from __future__ import annotations

import json
import re
import sys
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

REPO_ROOT = Path(__file__).resolve().parents[2]
sys.path.insert(0, str(REPO_ROOT))

import docx_header
from config import TROY_EMAIL, TROY_LINKEDIN, TROY_PHONE, TROY_PORTFOLIO
from docx_header import add_paragraph_bottom_border, build_navy_header


APP_DIR = Path(__file__).resolve().parent
OUTPUT_DIR = APP_DIR / "output"
BUILD_LOG_DIR = APP_DIR / "build_logs"
FONT = "EB Garamond"
BLACK = RGBColor(0x14, 0x14, 0x14)
GRAY = RGBColor(0x55, 0x55, 0x55)
STEEL = RGBColor(0x2D, 0x6A, 0x9F)


RESUME = {
    "summary": (
        "Senior investigator with 25 years of sworn law-enforcement service, including two investigative "
        "rotations and 5.5 years in electronic-crimes and digital-forensics assignments. Led complex financial, "
        "property, and digital investigations from complaint intake through evidence analysis, legal process, "
        "prosecutor referral, and courtroom support. Experienced in financial-transaction review, Excel-based "
        "evidence summaries, pattern recognition, interviews, multi-agency coordination, and clear, defensible "
        "reports. Certified Fraud Examiner credential actively being pursued through ACFE."
    ),
    "capabilities": (
        "Complex fraud investigations | Financial-record analysis | Pattern recognition | Case triage and prioritization | "
        "Interviews and statements | Evidence preservation | Search warrants and subpoenas | Defensible reports | "
        "Prosecutor and agency coordination | Restitution support | Confidential information handling | Excel | Microsoft 365"
    ),
    "results": [
        "Led a multi-victim Business Email Compromise investigation involving shell companies and interstate wire transfers that resulted in a felony conviction, documented verified victim losses exceeding $360,000, and earned written recognition from an Assistant Dakota County Attorney.",
        "Investigated approximately $80,000 in unauthorized company-card charges, collected and analyzed receipts, built an Excel financial summary, and delivered a complete case package that supported a felony Theft by Swindle conviction and court-ordered restitution.",
        "Consistently self-initiated fraud follow-up through resolution during patrol assignments between investigative rotations, a pattern documented in a written commendation from a supervising sergeant.",
        "Personally processed 5,304 GB of digital evidence in 2020 across Lakeville and partner-agency matters, organizing complex records and explaining findings to investigators, supervisors, and prosecutors.",
    ],
    "page1_jobs": [
        {
            "title": "Real Estate Consultant",
            "dates": "June 2024 - March 2026",
            "employer": "eXp Realty / KW Select | South Metro Minnesota",
            "bullets": [
                "Managed client records, negotiations, inspections, financing, title, and closing details across $3.2 million in completed residential sales during the transition from law enforcement."
            ],
        },
        {
            "title": "Police Officer",
            "dates": "January 2022 - May 2024",
            "employer": "Lakeville Police Department | Lakeville, Minnesota",
            "bullets": [
                "Returned to patrol after a specialized assignment and managed incident review, interviews, evidence, written reports, referrals, and follow-up across concurrent calls and investigations until medical retirement."
            ],
        },
        {
            "title": "Detective / Digital Forensic Examiner",
            "dates": "June 2017 - December 2021",
            "employer": "Dakota County Electronic Crimes Task Force, assigned from Lakeville Police Department | Minnesota",
            "bullets": [
                "Served as Lakeville's representative and digital-forensics subject-matter resource in a ten-agency task force, managing examination priorities and coordinating with investigators, supervisors, prosecutors, and partner agencies.",
                "Examined financial, cloud, mobile-device, computer, and provider records to identify relationships, timelines, contradictions, and evidence requiring further legal process.",
                "Produced defensible forensic reports and investigative summaries, maintained evidence integrity and confidentiality, and supported legal review and courtroom proceedings.",
            ],
        },
        {
            "title": "Detective / Electronic Crimes Unit",
            "dates": "September 2016 - June 2017",
            "employer": "Lakeville Police Department | Lakeville, Minnesota",
            "bullets": [
                "Built an investigator resource containing preservation, subpoena, and search-warrant templates with provider reference information, improving consistency in evidence requests and case documentation."
            ],
        },
    ],
    "page2_jobs": [
        {
            "title": "Police Officer / Field Training Officer",
            "dates": "June 2011 - August 2016",
            "employer": "Lakeville Police Department | Lakeville, Minnesota",
            "bullets": [
                "Managed fraud, theft, identity-related, and other investigations while assigned to patrol, conducting follow-up, interviews, record collection, report writing, and prosecutor referral without close supervision.",
                "Trained and evaluated officers on investigation, evidence, policy, documentation, communication, and sound decision-making.",
            ],
        },
        {
            "title": "Police Officer / Investigator",
            "dates": "March 2010 - May 2011",
            "employer": "Lakeville Police Department | Lakeville, Minnesota",
            "bullets": [
                "Conducted financial-crime and property-crime investigations, including receipt analysis, transaction reconstruction, surveillance review, search warrants, suspect interviews, and prosecutor-ready case files."
            ],
        },
        {
            "title": "Police Officer / Field Training Officer",
            "dates": "November 1998 - February 2010",
            "employer": "Lakeville Police Department | Lakeville, Minnesota",
            "bullets": [
                "Built the investigative foundation for later detective assignments through detailed report writing, interviews, evidence collection, follow-up, training duties, and work across multiple agencies."
            ],
        },
        {
            "title": "Adjunct Faculty / Criminal Justice",
            "dates": "March 2007 - October 2025",
            "employer": "University of Phoenix | Remote, concurrent with sworn service",
            "bullets": [
                "Taught undergraduate Criminal Justice courses remotely for 18 years, explaining investigative, legal, ethical, and evidence concepts to adult learners and evaluating written work against clear standards.",
                "Received the Phoenix500 Faculty Excellence Award in 2020 and 2021 and a Faculty of the Year nomination in 2021.",
            ],
        },
        {
            "title": "U.S. Army",
            "dates": "8 years 3 months",
            "employer": "Reserve, Active Duty, and Minnesota Army National Guard | Honorably Discharged",
            "bullets": [],
        },
    ],
}


COVER_PARAGRAPHS = [
    (
        "The Senior Fraud Investigator opportunity shown on LinkedIn caught my attention because the work has a clear endpoint: "
        "determine what the records support, document the case so it can withstand scrutiny, and move it toward recovery, corrective "
        "action, or referral. That is familiar investigative ground. I spent 25 years in Minnesota law enforcement, including two "
        "investigative rotations and 5.5 years assigned to electronic-crimes and digital-forensics work. I managed cases from intake "
        "through interviews, record collection, legal process, evidence analysis, prosecutor review, and courtroom support."
    ),
    (
        "Among several complex financial investigations, I led a multi-victim Business Email Compromise case involving shell companies "
        "and interstate wire transfers. The investigation documented verified losses exceeding $360,000, resulted in a felony conviction, "
        "and produced written recognition from an Assistant Dakota County Attorney for the documentation and legal coordination. That work "
        "required following money across records, identifying relationships, testing the case theory against conflicting information, and "
        "presenting the result clearly enough for legal review."
    ),
    (
        "In a separate occupational fraud matter, I investigated approximately $80,000 in unauthorized personal charges made on an "
        "employer's company card. I collected and analyzed receipts, reconstructed the transactions in Excel, and organized the evidence "
        "into a complete case package for the county attorney. The case resulted in a felony Theft by Swindle conviction and court-ordered "
        "restitution. During patrol assignments between investigative rotations, I also continued working fraud cases through resolution "
        "rather than forwarding the initial report. A supervising sergeant documented that pattern in a written commendation."
    ),
    (
        "My digital-forensics work adds another useful dimension. As Lakeville Police Department's representative and digital-forensics "
        "resource in a ten-agency task force, I managed competing examination priorities and worked with investigators, supervisors, "
        "prosecutors, technical providers, and partner agencies. I personally processed 5,304 GB of evidence in 2020. The job was not simply "
        "extracting data. It was finding the records that mattered, preserving their integrity, explaining what they proved and what they did "
        "not, and writing a report another decision-maker could defend."
    ),
    (
        "I want to be direct about the transition. I have not worked inside a healthcare payer, and I do not yet have hands-on experience "
        "with claim extraction, reimbursement analysis, coding review, overpayment recovery, HIPAA, or OIG reporting. That is the central "
        "gap in my candidacy. Advize's model is still appealing because its investigators work alongside clinical and coding auditors who "
        "already know the healthcare side deeply. I can contribute disciplined investigation, interviews, evidence handling, and defensible "
        "reporting while building payer-specific judgment. Fraud and improper billing shift costs to patients, plan members, employers, "
        "premium-payers, and taxpayers. That consequence matters, and so does getting the case right."
    ),
    (
        "Thank you for considering whether my investigative foundation can contribute to Advize's SIU work. I would welcome a candid "
        "conversation about the payer-side learning curve and the areas where I could add value immediately."
    ),
]


def configure_contact_parts() -> None:
    phone_digits = re.sub(r"[^0-9+]", "", TROY_PHONE)
    linkedin_url = TROY_LINKEDIN if TROY_LINKEDIN.startswith("http") else f"https://www.{TROY_LINKEDIN}"
    docx_header.CONTACT_PARTS = [
        *([(TROY_PHONE, f"tel:{phone_digits}")] if TROY_PHONE else []),
        (TROY_EMAIL, f"mailto:{TROY_EMAIL}"),
        (TROY_LINKEDIN, linkedin_url),
        ("troyhokanson.com", TROY_PORTFOLIO),
    ]


def set_run_font(run, size: float, *, bold: bool = False, italic: bool = False, color=BLACK) -> None:
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        r_fonts.set(qn(f"w:{attr}"), FONT)


def set_paragraph(paragraph, *, before=0, after=0, line=1.05, keep_next=False, keep_together=False) -> None:
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    pf.widow_control = True
    pf.keep_with_next = keep_next
    pf.keep_together = keep_together


def prepare_document(doc_type: str) -> Document:
    configure_contact_parts()
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)

    if doc_type == "resume":
        left = right = 0.65
        bottom = 0.55
        top = 1.46
        body_size = 10.25
    else:
        left = right = 0.78
        bottom = 0.68
        top = 1.52
        body_size = 10.5

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(body_size)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.05

    list_style = doc.styles["List Bullet"]
    list_style.font.name = FONT
    list_style.font.size = Pt(body_size)

    docx_header.NAME_FONT = FONT
    docx_header.CONTACT_FONT = FONT
    docx_header.BODY_FONT = FONT
    build_navy_header(
        doc,
        body_top_margin_inches=top,
        body_bottom_margin_inches=bottom,
        body_left_margin_inches=left,
        body_right_margin_inches=right,
    )
    return doc


def add_section_heading(doc: Document, text: str, *, page_break_before: bool = False) -> None:
    p = doc.add_paragraph()
    set_paragraph(p, before=8, after=4, line=1.0, keep_next=True)
    p.paragraph_format.page_break_before = page_break_before
    run = p.add_run(text.upper())
    set_run_font(run, 11, bold=True, color=STEEL)
    add_paragraph_bottom_border(p, color_hex="C9A84C", size=6)


def add_body_paragraph(doc: Document, text: str, *, size=10.25, after=4) -> None:
    p = doc.add_paragraph()
    set_paragraph(p, before=0, after=after, line=1.05, keep_together=True)
    set_run_font(p.add_run(text), size)


def add_bullet(doc: Document, text: str, *, size=10.25) -> None:
    p = doc.add_paragraph(style="List Bullet")
    set_paragraph(p, before=0, after=2, line=1.05, keep_together=True)
    p.paragraph_format.left_indent = Inches(0.22)
    p.paragraph_format.first_line_indent = Inches(-0.14)
    for run in p.runs:
        run.text = ""
    set_run_font(p.add_run(text), size)


def add_job(doc: Document, job: dict) -> None:
    p = doc.add_paragraph()
    set_paragraph(p, before=5, after=0, line=1.0, keep_next=True, keep_together=True)
    set_run_font(p.add_run(job["title"]), 10.5, bold=True)
    set_run_font(p.add_run(" | " + job["dates"]), 10.0, bold=True, color=GRAY)

    p = doc.add_paragraph()
    set_paragraph(p, before=0, after=2, line=1.0, keep_next=bool(job["bullets"]), keep_together=True)
    set_run_font(p.add_run(job["employer"]), 9.75, italic=True, color=GRAY)

    for bullet in job["bullets"]:
        add_bullet(doc, bullet)


def add_degree(doc: Document, lines: list[str]) -> None:
    p = doc.add_paragraph()
    set_paragraph(p, before=4, after=1, line=1.0, keep_together=True)
    for index, line in enumerate(lines):
        run = p.add_run(line)
        set_run_font(run, 9.75, bold=index == 0)
        if index < len(lines) - 1:
            run.add_break()


def build_resume() -> Path:
    doc = prepare_document("resume")
    add_section_heading(doc, "Professional Summary")
    add_body_paragraph(doc, RESUME["summary"], after=4)
    add_section_heading(doc, "Core Investigation Capabilities")
    add_body_paragraph(doc, RESUME["capabilities"], after=4)
    add_section_heading(doc, "Selected Fraud Investigation Results")
    for bullet in RESUME["results"]:
        add_bullet(doc, bullet)
    add_section_heading(doc, "Professional Experience")
    for job in RESUME["page1_jobs"]:
        add_job(doc, job)

    add_section_heading(doc, "Additional Experience", page_break_before=True)
    for job in RESUME["page2_jobs"]:
        add_job(doc, job)

    add_section_heading(doc, "Education")
    for degree in [
        ["Master of Arts, Police Leadership, Administration and Education", "University of St. Thomas, St. Paul, MN", "GPA: 3.94", "2005"],
        ["Bachelor of Arts, Criminal Justice, Magna Cum Laude", "St. Cloud State University, St. Cloud, MN", "GPA: 3.51", "1998"],
        ["Associate of Arts, Criminal Justice, Magna Cum Laude", "St. Cloud State University, St. Cloud, MN", "GPA: 3.50", "1996"],
    ]:
        add_degree(doc, degree)

    add_section_heading(doc, "Training and Certifications")
    for bullet in [
        "Certified Fraud Examiner (CFE), actively pursuing through ACFE, 2026.",
        "Certified Cyber Crime Investigator (CCCI) No. 4793, January 2023.",
        "TCORCA, Show Me the Money: Forensic Accounting Seminar, May 2017.",
        "Reid Technique of Interviewing and Interrogation.",
        "BCA Law Enforcement Supervision and Management Program, 98 hours.",
    ]:
        add_bullet(doc, bullet, size=9.75)

    output = OUTPUT_DIR / "Hokanson_Resume_Advize_Senior_Fraud_Investigator.docx"
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    sanitize_docx_fonts(output)
    return output


def add_cover_line(doc: Document, text: str, *, after=0, bold=False) -> None:
    p = doc.add_paragraph()
    set_paragraph(p, before=0, after=after, line=1.0, keep_together=True)
    set_run_font(p.add_run(text), 10.5, bold=bold)


def build_cover() -> Path:
    doc = prepare_document("cover")
    add_cover_line(doc, "August 20, 2026", after=8)
    add_cover_line(doc, "Hiring Manager")
    add_cover_line(doc, "Advize Health")
    add_cover_line(doc, "Orlando, Florida", after=9)
    add_cover_line(doc, "Dear Hiring Manager,", after=9)

    for paragraph in COVER_PARAGRAPHS:
        p = doc.add_paragraph()
        set_paragraph(p, before=0, after=8, line=1.05, keep_together=True)
        set_run_font(p.add_run(paragraph), 10.5)

    p = doc.add_paragraph()
    set_paragraph(p, before=5, after=0, line=1.0, keep_next=True)
    set_run_font(p.add_run("Respectfully,"), 10.5)
    p = doc.add_paragraph()
    set_paragraph(p, before=38, after=0, line=1.0, keep_together=True)
    set_run_font(p.add_run("Troy Hokanson"), 10.5)

    output = OUTPUT_DIR / "Hokanson_Cover_Advize_Senior_Fraud_Investigator.docx"
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    sanitize_docx_fonts(output)
    return output


def sanitize_docx_fonts(path: Path) -> None:
    replacements = {
        "Calibri Light": FONT,
        "Calibri": FONT,
        "Aptos Display": FONT,
        "Aptos": FONT,
        "Arial": FONT,
    }
    temporary = path.with_suffix(".fonts.docx")
    with zipfile.ZipFile(path, "r") as source, zipfile.ZipFile(temporary, "w", zipfile.ZIP_DEFLATED) as target:
        for item in source.infolist():
            data = source.read(item.filename)
            if item.filename.startswith("word/") and item.filename.endswith((".xml", ".rels")):
                text = data.decode("utf-8", errors="replace")
                for old, new in replacements.items():
                    text = re.sub(re.escape(old), new, text, flags=re.IGNORECASE)
                data = text.encode("utf-8")
            target.writestr(item, data)
    temporary.replace(path)


def inspect_docx(path: Path, doc_type: str) -> dict:
    doc = Document(path)
    section = doc.sections[0]
    expected = {
        "resume": {"left": 0.65, "right": 0.65, "bottom": 0.55, "top": 1.46},
        "cover": {"left": 0.78, "right": 0.78, "bottom": 0.68, "top": 1.52},
    }[doc_type]
    bullet_paragraphs = [p for p in doc.paragraphs if p.style and p.style.name == "List Bullet"]
    keep_next = sum(1 for p in doc.paragraphs if p.paragraph_format.keep_with_next)
    with zipfile.ZipFile(path) as archive:
        xml = "\n".join(
            archive.read(name).decode("utf-8", errors="replace")
            for name in archive.namelist()
            if name.startswith("word/") and name.endswith(".xml")
        ).lower()
        rels = "\n".join(
            archive.read(name).decode("utf-8", errors="replace")
            for name in archive.namelist()
            if name.startswith("word/") and name.endswith(".rels")
        )
    actual = {
        "left": round(section.left_margin.inches, 3),
        "right": round(section.right_margin.inches, 3),
        "bottom": round(section.bottom_margin.inches, 3),
        "top": round(section.top_margin.inches, 3),
    }
    checks = {
        "us_letter": round(section.page_width.inches, 2) == 8.5 and round(section.page_height.inches, 2) == 11,
        "margins_match": actual == expected,
        "header_part_present": bool(doc.sections[0].header.paragraphs),
        "real_list_style_present": len(bullet_paragraphs) > 0 if doc_type == "resume" else True,
        "keep_with_next_present": keep_next > 0,
        "garamond_present": "eb garamond" in xml,
        "calibri_absent": "calibri" not in xml,
        "aptos_absent": "aptos" not in xml,
        "arial_absent": "arial" not in xml,
        "contact_hyperlinks_present": "hyperlink" in rels.lower(),
        "city_omitted_from_header": "lakeville, mn" not in " ".join(p.text.lower() for p in doc.sections[0].header.paragraphs),
    }
    return {
        "path": str(path),
        "document_type": doc_type,
        "paragraph_count": len(doc.paragraphs),
        "bullet_paragraph_count": len(bullet_paragraphs),
        "keep_with_next_count": keep_next,
        "margins_inches": actual,
        "checks": checks,
        "passed": all(checks.values()),
    }


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    BUILD_LOG_DIR.mkdir(parents=True, exist_ok=True)
    resume = build_resume()
    cover = build_cover()
    report = {
        "layout_contract": str(APP_DIR / "layout_contract.json"),
        "documents": [inspect_docx(resume, "resume"), inspect_docx(cover, "cover")],
    }
    report["passed"] = all(item["passed"] for item in report["documents"])
    (BUILD_LOG_DIR / "docx_structure_audit.json").write_text(json.dumps(report, indent=2), encoding="utf-8")
    print(json.dumps(report, indent=2))
    if not report["passed"]:
        raise SystemExit(1)


if __name__ == "__main__":
    main()

