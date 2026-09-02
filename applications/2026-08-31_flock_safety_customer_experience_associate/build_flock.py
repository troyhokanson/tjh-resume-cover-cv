#!/usr/bin/env python3

from __future__ import annotations

import hashlib
import json
import re
import sys
import zipfile
from copy import deepcopy
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


APP_DIR = Path(__file__).resolve().parent
REPO_ROOT = APP_DIR.parents[1]
OUTPUT_DIR = APP_DIR / "output"
BUILD_LOG_DIR = APP_DIR / "build_logs"
sys.path.insert(0, str(REPO_ROOT))

from docx_header import (  # noqa: E402
    BLACK,
    GRAY,
    STEEL,
    add_bullet,
    add_job_block,
    add_section_heading,
    build_navy_header,
    new_document,
    set_paragraph_format,
    set_run,
)


FONT = "EB Garamond"
BUILD_DATE = "2026-08-31"
REQUISITION = "1cabec78-f1bd-4615-95d2-5d8196eb46e0"
ROLE = "Customer Experience Associate"
EMPLOYER = "Flock Safety"


RESUME_JOBS_PAGE_1 = [
    {
        "title": "Independent Professional",
        "employer": "Independent | Remote",
        "dates": "April 2026 - Present",
        "bullets": [
            "Pursuing structured professional development in fraud examination, SQL, Python, and technology-supported investigative workflows while targeting hands-on public-safety technology and customer-support work.",
        ],
    },
    {
        "title": "Real Estate Consultant",
        "employer": "eXp Realty / KW Select | South Metro MN",
        "dates": "June 2024 - June 2026",
        "bullets": [
            "Managed client relationships from initial consultation through negotiation, inspection, financing, title, and closing, completing $3.2 million in residential sales during the transition from law enforcement.",
            "Coordinated clients, lenders, inspectors, appraisers, title professionals, and other agents by phone, email, and in person to explain options, resolve barriers, and keep time-sensitive transactions moving.",
        ],
    },
    {
        "title": "Police Officer",
        "employer": "Lakeville Police Department | Lakeville, MN",
        "dates": "January 2022 - May 2024",
        "bullets": [
            "Used Axon Body 3, Axon Fleet 3, Axon Evidence, Motorola radios, and Microsoft Office in daily patrol, documentation, and evidence workflows.",
            "Handled urgent public inquiries, explained procedures to people under stress, and maintained accurate reports and follow-up across changing priorities.",
        ],
    },
    {
        "title": "Detective / Digital Forensic Examiner",
        "employer": "Dakota County Electronic Crimes Task Force, assigned from Lakeville Police Department | Lakeville, MN",
        "dates": "June 2017 - December 2021",
        "bullets": [
            "Served as Lakeville's digital-forensics resource in a 10-agency task-force environment, supporting officers, investigators, supervisors, prosecutors, and partner agencies with device, evidence, and workflow questions.",
            "Resolved a commercial burglary by combining surveillance, physical evidence, provider preservation, cloud legal process, and forensic analysis. The case produced a felony conviction and written recognition for persistence and for strengthening a local business's confidence in the department.",
            "Managed concurrent examinations and investigations, documented technical findings for non-technical decision-makers, and maintained clear follow-up when work required assistance from other agencies or service providers.",
        ],
    },
]


RESUME_JOBS_PAGE_2 = [
    {
        "title": "Detective / Electronic Crimes Unit",
        "employer": "Lakeville Police Department | Lakeville, MN",
        "dates": "September 2016 - June 2017",
        "bullets": [
            "Acquired and configured the unit's initial Cellebrite UFED and supported investigators troubleshooting mobile-device and digital-evidence questions.",
            "Built a reusable investigator resource library with preservation-request, administrative-subpoena, search-warrant, and service-provider guidance to support consistent electronic-evidence work.",
        ],
    },
    {
        "title": "Police Officer / Field Training Officer",
        "employer": "Lakeville Police Department | Lakeville, MN",
        "dates": "June 2011 - August 2016",
        "bullets": [
            "Returned to patrol after the first investigative rotation, training and coaching officers and supporting reserve and park-ranger development.",
        ],
    },
    {
        "title": "Police Officer / Investigator (First Investigative Rotation)",
        "employer": "Lakeville Police Department | Lakeville, MN",
        "dates": "March 2010 - May 2011",
        "bullets": [
            "Managed financial and property cases from intake through interviews, records review, legal process, documentation, and prosecutorial handoff during a full-time investigative rotation.",
        ],
    },
    {
        "title": "Police Officer / Field Training Officer",
        "employer": "Lakeville Police Department | Lakeville, MN",
        "dates": "November 1998 - February 2010",
        "bullets": [
            "Led agency-side operations for a $40,000 Target-funded Genetec AutoVu ALPR project from 2007 to 2010, coordinating Target, department leadership, city IT, BCA CJIS, Genetec, nightly hotlist synchronization, and field troubleshooting.",
            "Received 20+ written commendations across the law-enforcement career.",
        ],
    },
    {
        "title": "Adjunct Faculty / Criminal Justice",
        "employer": "University of Phoenix | Remote, concurrent with sworn service",
        "dates": "March 2007 - October 2025",
        "bullets": [
            "Taught undergraduate Criminal Justice courses remotely for 18 years, explaining investigative, legal, and technical concepts through structured lessons, written feedback, and individual learner support.",
            "Received Phoenix500 Faculty Excellence Awards in 2020 and 2021 and a Faculty of the Year nomination in 2021.",
        ],
    },
]


COVER_PARAGRAPHS = [
    (
        "Flock's Customer Experience Associate role is not passive or scripted support. The job calls for someone who can "
        "investigate a technical problem, explain the answer clearly, keep ownership across teams, and improve the next "
        "customer's path through better documentation. That combination is familiar from my work with public-safety technology. "
        "I want to bring that same discipline to Flock in a customer-facing role where clear answers and steady follow-through matter every day, even when priorities change quickly."
    ),
    (
        "One of my strongest examples is a $40,000 Target-funded Genetec AutoVu license-plate-recognition project. I coordinated "
        "the operational use case with Target, department leadership, Lakeville IT, BCA CJIS, Genetec, and the officers using the "
        "system. The work included a nightly data-synchronization process and troubleshooting cameras and connectivity. It taught me "
        "that public-safety technology earns trust one resolved issue at a time, especially under real operational pressure."
    ),
    (
        "I later acquired and configured Lakeville's initial Cellebrite UFED and built a structured investigator resource library "
        "with preservation, subpoena, search-warrant, and service-provider guidance. Officers and investigators brought me device and "
        "digital-evidence questions, and I kept the answers usable across different levels of technical experience. A 2020 performance review put it plainly: \"With the ECU training, Troy has become "
        "a subject matter expert with digital forensics and is a resource to his peers with electronic crime investigations.\""
    ),
    (
        "The same approach carried into case work. In a commercial burglary investigation, I connected surveillance, physical evidence, "
        "provider preservation, cloud legal process, and forensic analysis into a defensible timeline. The case resulted in a felony conviction. "
        "Written supervisory recognition credited the work with strengthening the local business's confidence in the department. Throughout my "
        "career, the person waiting for an answer has been part of the work."
    ),
    (
        "My background is not a traditional SaaS-support path. I have not held a commercial support title, managed SLA or ticket-volume metrics, "
        "or worked in Jira. What transfers directly is technical troubleshooting, concurrent case ownership, escalation, knowledge-resource development, "
        "calm customer communication, and disciplined documentation. Eighteen years of remote college instruction and recent client work in $3.2 million "
        "of residential transactions show that I can explain complicated information, listen carefully, and follow through."
    ),
    (
        "I am intentionally pursuing a hands-on individual-contributor role where I can learn the platform, resolve the issue in front of the customer, and improve "
        "the support path behind it. I would value the opportunity to discuss how my public-safety operator perspective could contribute to Flock's team."
    ),
]


def sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def set_keep(paragraph, *, keep_next=False, keep_together=False, widow=True) -> None:
    paragraph.paragraph_format.keep_with_next = keep_next
    paragraph.paragraph_format.keep_together = keep_together
    p_pr = paragraph._p.get_or_add_pPr()
    if widow and p_pr.find(qn("w:widowControl")) is None:
        p_pr.append(OxmlElement("w:widowControl"))


def configure_styles(doc: Document) -> None:
    for style in doc.styles:
        if style.type == 1:
            style.font.name = FONT
            style._element.get_or_add_rPr()
            r_fonts = style._element.rPr.rFonts
            if r_fonts is None:
                r_fonts = OxmlElement("w:rFonts")
                style._element.rPr.insert(0, r_fonts)
            for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
                r_fonts.set(qn(f"w:{attr}"), FONT)
    normal = doc.styles["Normal"]
    normal.font.size = Pt(10.25)
    list_style = doc.styles["List Bullet"]
    list_style.font.size = Pt(10.25)


def prepare_document(doc_type: str) -> Document:
    doc = new_document()
    configure_styles(doc)
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    if doc_type == "resume":
        build_navy_header(
            doc,
            body_top_margin_inches=1.46,
            body_bottom_margin_inches=0.55,
            body_left_margin_inches=0.65,
            body_right_margin_inches=0.65,
        )
    else:
        build_navy_header(
            doc,
            body_top_margin_inches=1.52,
            body_bottom_margin_inches=0.68,
            body_left_margin_inches=0.78,
            body_right_margin_inches=0.78,
        )
    return doc


def set_metadata(doc: Document, document_type: str) -> None:
    props = doc.core_properties
    label = "Resume" if document_type == "resume" else "Cover Letter"
    props.author = "Troy Hokanson"
    props.last_modified_by = "Troy Hokanson"
    props.title = f"Troy Hokanson - {EMPLOYER} - {ROLE} - {REQUISITION} - {label} - {BUILD_DATE}"
    props.subject = f"Application for {EMPLOYER} {ROLE}, requisition {REQUISITION}"
    props.keywords = f"Troy Hokanson, {EMPLOYER}, {ROLE}, {REQUISITION}, {label}, {BUILD_DATE}"
    props.comments = ""
    timestamp = datetime(2026, 8, 31, 22, 0, tzinfo=timezone.utc)
    props.created = timestamp
    props.modified = timestamp


def add_headline(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(p, before=0, after=4, line=1.0)
    set_keep(p, keep_next=True, keep_together=True)
    set_run(p.add_run(text), font=FONT, size=11, bold=True, color=STEEL)


def add_body_paragraph(doc: Document, text: str, *, size=10.25, after=4, line=1.05, keep=False) -> None:
    p = doc.add_paragraph()
    set_paragraph_format(p, before=0, after=after, line=line)
    set_keep(p, keep_together=keep)
    set_run(p.add_run(text), font=FONT, size=size, color=BLACK)


def add_compact_line(doc: Document, text: str, *, bold=False, after=2) -> None:
    p = doc.add_paragraph()
    set_paragraph_format(p, before=0, after=after, line=1.0)
    set_keep(p, keep_together=True)
    set_run(p.add_run(text), font=FONT, size=9.75, bold=bold, color=BLACK)


def add_section(doc: Document, text: str) -> None:
    p = add_section_heading(doc, text)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.05
    set_keep(p, keep_next=True, keep_together=True)
    for run in p.runs:
        set_run(run, font=FONT, size=11, bold=True, color=STEEL)


def add_job(doc: Document, job: dict[str, object]) -> None:
    title_p, employer_p = add_job_block(doc, str(job["title"]), str(job["employer"]), str(job["dates"]))
    title_p.paragraph_format.space_before = Pt(5)
    title_p.paragraph_format.space_after = Pt(0)
    title_p.paragraph_format.line_spacing = 1.05
    employer_p.paragraph_format.space_before = Pt(0)
    employer_p.paragraph_format.space_after = Pt(2)
    employer_p.paragraph_format.line_spacing = 1.05
    set_keep(title_p, keep_next=True, keep_together=True)
    set_keep(employer_p, keep_next=True, keep_together=True)
    for run in title_p.runs:
        set_run(run, font=FONT, size=10.5, bold=True, color=BLACK)
    for run in employer_p.runs:
        set_run(run, font=FONT, size=9.75, italic=run.italic, color=GRAY)
    bullets = list(job["bullets"])
    for index, bullet in enumerate(bullets):
        p = add_bullet(doc, str(bullet), size=10.25)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.05
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        set_keep(p, keep_together=index == 0)


def add_degree(doc: Document, lines: list[str]) -> None:
    p = doc.add_paragraph()
    set_paragraph_format(p, before=1, after=2, line=1.0)
    set_keep(p, keep_together=True)
    for index, line in enumerate(lines):
        run = p.add_run(line)
        set_run(run, font=FONT, size=9.35, bold=index == 0, color=BLACK if index == 0 else GRAY)
        if index < len(lines) - 1:
            run.add_break()


def sanitize_non_header_fonts(path: Path) -> None:
    replacements = {
        "Calibri Light": FONT,
        "Calibri": FONT,
        "Aptos Display": FONT,
        "Aptos": FONT,
        "Arial": FONT,
    }
    temporary = path.with_suffix(".font-clean.docx")
    with zipfile.ZipFile(path, "r") as source, zipfile.ZipFile(temporary, "w", zipfile.ZIP_DEFLATED) as target:
        for item in source.infolist():
            data = source.read(item.filename)
            is_header_part = item.filename.startswith("word/header")
            if item.filename.startswith("word/") and item.filename.endswith((".xml", ".rels")) and not is_header_part:
                text = data.decode("utf-8", errors="replace")
                for old, new in replacements.items():
                    text = re.sub(re.escape(old), new, text, flags=re.IGNORECASE)
                data = text.encode("utf-8")
            target.writestr(item, data)
    temporary.replace(path)


def synchronize_alternate_header_parts(path: Path) -> None:
    temporary = path.with_suffix(".alternate-headers.docx")
    with zipfile.ZipFile(path, "r") as source:
        payloads = {item.filename: source.read(item.filename) for item in source.infolist()}
        infos = {item.filename: item for item in source.infolist()}

    rels_name = "word/_rels/document.xml.rels"
    rels_text = payloads[rels_name].decode("utf-8")
    document_name = "word/document.xml"
    document_text = payloads[document_name].decode("utf-8")
    header_refs = dict(re.findall(r'<w:headerReference w:type="([^"]+)" r:id="([^"]+)"/>', document_text))
    relationship_targets = dict(
        re.findall(
            r'<Relationship Id="([^"]+)" Type="[^"]+/header" Target="([^"]+)"/>',
            rels_text,
        )
    )
    required = {"default", "even", "first"}
    if not required.issubset(header_refs):
        raise RuntimeError(f"Missing header references: {sorted(required - set(header_refs))}")
    default_target = relationship_targets[header_refs["default"]]
    default_part = f"word/{default_target}"
    default_rels = f"word/_rels/{default_target}.rels"
    for header_type in ("even", "first"):
        target = relationship_targets[header_refs[header_type]]
        target_part = f"word/{target}"
        target_rels = f"word/_rels/{target}.rels"
        payloads[target_part] = payloads[default_part]
        payloads[target_rels] = payloads[default_rels]

    with zipfile.ZipFile(temporary, "w", zipfile.ZIP_DEFLATED) as target:
        for name, data in payloads.items():
            target.writestr(infos[name] if name in infos else name, data)
    temporary.replace(path)


def build_resume() -> Path:
    doc = prepare_document("resume")
    add_headline(doc, "Public Safety Technology Support | Technical Troubleshooting | Customer Communication")
    add_section(doc, "Professional Summary")
    add_body_paragraph(
        doc,
        "Former detective and digital forensic examiner with 25 years of public-safety service, 18 years of remote college instruction, and recent private-sector client work. Supported officers, investigators, supervisors, students, and clients through technical questions, time-sensitive problems, and complex documentation. Known for case ownership, practical explanations, calm communication, and follow-through.",
        after=4,
    )
    add_section(doc, "Core Qualifications")
    add_body_paragraph(
        doc,
        "Technical troubleshooting | Customer and stakeholder communication | Hardware and software workflows | Cross-functional escalation | Concurrent case prioritization | Knowledge resources and training documentation | Phone and email support | Public-safety operations | Digital evidence | Microsoft Office",
        after=2,
    )
    add_section(doc, "Professional Experience")
    for job in RESUME_JOBS_PAGE_1:
        add_job(doc, job)

    section = doc.sections[0]
    doc.settings.odd_and_even_pages_header_footer = True
    section.different_first_page_header_footer = True
    section.even_page_header.is_linked_to_previous = False
    section.first_page_header.is_linked_to_previous = False
    _ = section.even_page_header.part
    _ = section.first_page_header.part

    for job in RESUME_JOBS_PAGE_2:
        add_job(doc, job)

    add_section(doc, "Military Service")
    add_compact_line(
        doc,
        "U.S. Army, 8 years 3 months. Reserve, Active Duty, and Minnesota Army National Guard. Infantry, Armor, Motor Transport, and Military Police. Honorably Discharged.",
    )
    add_section(doc, "Technology, Training, and Credentials")
    for line in [
        "Public-safety technology: Genetec AutoVu, Axon Body 3, Axon Fleet 3, Axon Evidence, Motorola radios, Cellebrite Touch 2, UFED 4PC, Physical Analyzer, GrayKey, Magnet AXIOM, X-Ways Forensics, and Microsoft Office.",
        "LexisNexis Accurint, 6.5 years of direct investigative use across both investigative rotations.",
        "Certified Cyber Crime Investigator (CCCI) No. 4793, January 2023.",
        "BCA Law Enforcement Supervision & Management Program, 98 hours, 2012. University of Phoenix Certified Advanced Facilitator, 2012.",
    ]:
        add_compact_line(doc, line, after=1)

    add_section(doc, "Education")
    add_degree(doc, [
        "Master of Arts, Police Leadership, Administration and Education | University of St. Thomas, St. Paul, MN | GPA 3.94 | 2005",
    ])
    add_degree(doc, [
        "Bachelor of Arts, Criminal Justice, Magna Cum Laude | St. Cloud State University, St. Cloud, MN | GPA 3.51 | 1998",
    ])
    add_degree(doc, [
        "Associate of Arts, Criminal Justice, Magna Cum Laude | St. Cloud State University, St. Cloud, MN | GPA 3.50 | 1996",
    ])

    set_metadata(doc, "resume")
    output = OUTPUT_DIR / "2026-08-31_Troy-Hokanson_Flock-Customer-Experience-Associate_1cabec78_Resume.docx"
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    sanitize_non_header_fonts(output)
    synchronize_alternate_header_parts(output)
    return output


def add_cover_line(doc: Document, text: str, *, after=0, bold=False) -> None:
    p = doc.add_paragraph()
    set_paragraph_format(p, before=0, after=after, line=1.0)
    set_keep(p, keep_together=True)
    set_run(p.add_run(text), font=FONT, size=10.25, bold=bold, color=BLACK)


def build_cover() -> Path:
    doc = prepare_document("cover")
    add_cover_line(doc, "August 31, 2026", after=4)
    add_cover_line(doc, "Hiring Manager | Flock Safety | Atlanta, GA 30327", after=5)
    add_cover_line(doc, "Dear Hiring Manager,", after=5)

    for paragraph in COVER_PARAGRAPHS:
        p = doc.add_paragraph()
        set_paragraph_format(p, before=0, after=4, line=1.05)
        set_keep(p, keep_together=True)
        set_run(p.add_run(paragraph), font=FONT, size=10.25, color=BLACK)

    p = doc.add_paragraph()
    set_paragraph_format(p, before=0, after=0, line=1.0)
    set_keep(p, keep_next=True, keep_together=True)
    set_run(p.add_run("Respectfully,"), font=FONT, size=10.25, color=BLACK)
    p = doc.add_paragraph()
    set_paragraph_format(p, before=38, after=0, line=1.0)
    set_keep(p, keep_together=True)
    set_run(p.add_run("Troy Hokanson"), font=FONT, size=10.25, color=BLACK)

    set_metadata(doc, "cover")
    output = OUTPUT_DIR / "2026-08-31_Troy-Hokanson_Flock-Customer-Experience-Associate_1cabec78_Cover-Letter.docx"
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    sanitize_non_header_fonts(output)
    return output


def inspect_docx(path: Path, doc_type: str) -> dict[str, object]:
    doc = Document(path)
    section = doc.sections[0]
    expected = {
        "resume": {"left": 0.65, "right": 0.65, "bottom": 0.55, "top": 1.46},
        "cover": {"left": 0.78, "right": 0.78, "bottom": 0.68, "top": 1.52},
    }[doc_type]
    with zipfile.ZipFile(path) as archive:
        names = archive.namelist()
        xml = "\n".join(
            archive.read(name).decode("utf-8", errors="replace")
            for name in names
            if name.startswith("word/") and name.endswith(".xml")
        ).lower()
        document_xml = archive.read("word/document.xml").decode("utf-8", errors="replace")
        relationship_xml = "\n".join(
            archive.read(name).decode("utf-8", errors="replace")
            for name in names
            if name.startswith("word/_rels/") and name.endswith(".rels")
        )
        header_names = [name for name in names if name.startswith("word/header") and name.endswith(".xml")]
        header_xml = "\n".join(archive.read(name).decode("utf-8", errors="replace") for name in header_names).lower()
    actual = {
        "left": round(section.left_margin.inches, 2),
        "right": round(section.right_margin.inches, 2),
        "bottom": round(section.bottom_margin.inches, 2),
        "top": round(section.top_margin.inches, 2),
    }
    bullet_paragraphs = [p for p in doc.paragraphs if p.style and p.style.name == "List Bullet"]
    checks = {
        "us_letter": round(section.page_width.inches, 2) == 8.5 and round(section.page_height.inches, 2) == 11,
        "margins_match": actual == expected,
        "header_reference_present": "headerReference" in document_xml,
        "header_part_present": bool(header_names),
        "locked_header_background_marker": "troylockedheadernavybackground" in header_xml,
        "navy_and_gold_in_header": "0d1b2a" in header_xml and "c9a84c" in header_xml,
        "header_not_mutated_by_font_cleanup": "eb garamond" in header_xml and "troy hokanson" in header_xml,
        "real_bullet_style_present": len(bullet_paragraphs) > 0 if doc_type == "resume" else True,
        "keep_with_next_present": xml.count("w:keepnext") > 0,
        "garamond_present": "eb garamond" in xml,
        "calibri_absent": "calibri" not in xml,
        "aptos_absent": "aptos" not in xml,
        "arial_absent": "arial" not in xml,
        "contact_hyperlinks_present": relationship_xml.lower().count("hyperlink") >= 4,
        "no_comments_part": not any("comments" in name.lower() for name in names),
        "no_tracked_changes": "<w:ins" not in document_xml and "<w:del" not in document_xml,
        "no_hidden_text": "<w:vanish" not in xml,
    }
    return {
        "path": str(path),
        "sha256": sha256(path),
        "document_type": doc_type,
        "section_count": len(doc.sections),
        "paragraph_count": len(doc.paragraphs),
        "bullet_paragraph_count": len(bullet_paragraphs),
        "margins_inches": actual,
        "checks": checks,
        "passed": all(checks.values()),
    }


def main() -> int:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    BUILD_LOG_DIR.mkdir(parents=True, exist_ok=True)
    resume = build_resume()
    cover = build_cover()
    audits = {
        "resume": inspect_docx(resume, "resume"),
        "cover_letter": inspect_docx(cover, "cover"),
    }
    (BUILD_LOG_DIR / "docx_structure_audit.json").write_text(json.dumps(audits, indent=2), encoding="utf-8")
    print(json.dumps({
        "resume": str(resume),
        "cover_letter": str(cover),
        "audits_passed": all(item["passed"] for item in audits.values()),
    }, indent=2))
    return 0 if all(item["passed"] for item in audits.values()) else 1


if __name__ == "__main__":
    raise SystemExit(main())
