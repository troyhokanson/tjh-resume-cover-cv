#!/usr/bin/env python3
"""Build Troy Hokanson's Hugging Face Wild Card application documents."""

from pathlib import Path
import sys
import zipfile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[2]
sys.path.insert(0, str(ROOT))

from config import TROY_NAME  # noqa: E402
from docx_header import build_navy_header  # noqa: E402


OUT = Path(__file__).resolve().parent / "output"
GARAMOND = "Garamond"
NAVY = "0D1B2A"
GOLD = "C9A84C"
STEEL = "2D6A9F"
BLACK = "141414"
GRAY = "555555"


def set_cell_border_bottom(paragraph, color=GOLD, size=6, space=1):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def style_run(run, size=12, bold=False, italic=False, color=BLACK):
    run.font.name = GARAMOND
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)
    r_pr = run._element.get_or_add_rPr()
    fonts = r_pr.find(qn("w:rFonts"))
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        r_pr.append(fonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{attr}"), GARAMOND)


def style_document(doc, body_size):
    for style in doc.styles:
        if hasattr(style, "font"):
            style.font.name = GARAMOND
            style.font.size = Pt(body_size)
            r_pr = style.element.get_or_add_rPr()
            fonts = r_pr.find(qn("w:rFonts"))
            if fonts is None:
                fonts = OxmlElement("w:rFonts")
                r_pr.append(fonts)
            for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
                fonts.set(qn(f"w:{attr}"), GARAMOND)
    normal = doc.styles["Normal"]
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.05
    normal.paragraph_format.widow_control = True


def set_para(paragraph, before=0, after=0, line=1.05, keep_next=False):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.keep_with_next = keep_next
    fmt.widow_control = True


def add_target_line(doc, text):
    p = doc.add_paragraph()
    set_para(p, before=18, after=6, line=1.0, keep_next=True)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text.upper())
    style_run(run, size=12, bold=True, color=STEEL)
    set_cell_border_bottom(p)
    return p


def add_section(doc, text, first=False):
    p = doc.add_paragraph()
    set_para(p, before=12 if first else 10, after=4, line=1.0, keep_next=True)
    run = p.add_run(text.upper())
    style_run(run, size=12, bold=True, color=STEEL)
    set_cell_border_bottom(p)
    return p


def add_body(doc, text, size=12, after=4):
    p = doc.add_paragraph()
    set_para(p, after=after, line=1.05)
    run = p.add_run(text)
    style_run(run, size=size)
    return p


def add_job(doc, title, employer, dates):
    title_p = doc.add_paragraph()
    set_para(title_p, before=6, after=1, line=1.0, keep_next=True)
    style_run(title_p.add_run(title), size=12, bold=True, color=GOLD)

    employer_p = doc.add_paragraph()
    set_para(employer_p, after=3, line=1.0, keep_next=True)
    style_run(employer_p.add_run(employer), size=12, italic=True)
    style_run(employer_p.add_run(" | " + dates), size=12, color=GRAY)
    return title_p, employer_p


def add_bullet(doc, text, size=12, keep_next=False):
    p = doc.add_paragraph(style="List Bullet")
    set_para(p, after=2, line=1.05, keep_next=keep_next)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    run = p.add_run(text)
    style_run(run, size=size)
    return p


def add_degree(doc, degree, school, details):
    p = doc.add_paragraph()
    set_para(p, before=5, after=1, line=1.0, keep_next=True)
    style_run(p.add_run(degree), size=12, bold=True, color=GOLD)
    p2 = doc.add_paragraph()
    set_para(p2, after=1, line=1.0)
    style_run(p2.add_run(f"{school} | {details}"), size=12)


def new_branded_document(body_size):
    doc = Document()
    build_navy_header(
        doc,
        body_top_margin_inches=1.55,
        body_bottom_margin_inches=0.55,
        body_left_margin_inches=0.60,
        body_right_margin_inches=0.60,
    )
    style_document(doc, body_size)
    return doc


def sanitize_ooxml_font_defaults(path):
    """Remove dormant default-font names that fail the repository font gate."""
    temp = path.with_suffix(".tmp.docx")
    with zipfile.ZipFile(path, "r") as source, zipfile.ZipFile(
        temp, "w", zipfile.ZIP_DEFLATED
    ) as target:
        for item in source.infolist():
            data = source.read(item.filename)
            if item.filename in {"word/fontTable.xml", "word/theme/theme1.xml"}:
                data = data.replace(b"Calibri", b"Garamond")
                data = data.replace(b"Arial", b"Garamond")
            target.writestr(item, data)
    temp.replace(path)


def build_resume(path):
    doc = new_branded_document(12)
    props = doc.core_properties
    props.title = "Troy Hokanson Resume, Hugging Face Wild Card"
    props.subject = "Investigations, AI-Assisted Workflows and Public-Sector Enablement"
    props.author = TROY_NAME

    add_target_line(doc, "Investigations, AI-Assisted Workflows and Public-Sector Enablement")

    add_section(doc, "Professional Summary")
    add_body(
        doc,
        "Senior investigator, digital forensic examiner, instructor and workflow builder with 25 years of public-safety experience and 18 years teaching Criminal Justice remotely. Turns incomplete digital, financial, public-record and interview evidence into clear findings for nontechnical decision-makers. Uses generative AI and source-controlled systems to organize evidence, build repeatable workflows, test outputs and document human-review gates. Offers a practitioner perspective on technology-enabled wrongdoing, public-sector needs and user education.",
    )

    add_section(doc, "Applied AI and Workflow Project")
    add_job(doc, "AI-Assisted Career Search and Evidence System", "Independent project", "2026")
    project_bullets = [
        "Designed and tested a source-controlled system connecting job discovery, eligibility rules, evidence selection, documents, validation and status tracking across GitHub, Notion, Google Drive and Telegram.",
        "Defined deterministic location, travel, compensation, fit, privacy and candidate-separation rules for a 16-source search pipeline. Reviews false positives, duplicate alerts, incomplete listings and unintended suppression through source-verification, privacy, human-approval, rendering and repeated quality-review gates.",
    ]
    for i, bullet in enumerate(project_bullets):
        add_bullet(doc, bullet, keep_next=i == 0)

    add_section(doc, "Professional Experience")
    roles = [
        (
            "Adjunct Faculty, Criminal Justice",
            "University of Phoenix, Remote, concurrent with sworn service",
            "March 2007 - October 2025",
            [
                "Taught investigations, ethics, leadership, policy and evidence concepts to adult learners for 18 years. Translated complex material into structured instruction, written feedback and practical decision frameworks.",
                "Earned Certified Advanced Facilitator status and Phoenix500 Faculty Excellence recognition in 2020 and 2021.",
            ],
        ),
        (
            "Police Officer / Detective / Digital Forensic Examiner / Field Training Officer",
            "Lakeville Police Department, Lakeville, MN",
            "November 1998 - May 2024",
            [
                "Connected device artifacts, cloud records, telecommunications data, financial records, OSINT, interviews and legal-process returns in complex technology-enabled investigations.",
                "Served as Lakeville Police Department's digital-forensics resource within a 10-agency task-force structure and completed examinations for Lakeville and partner agencies. Processed 5,304 GB of digital evidence in 2020.",
                "Wrote search-warrant affidavits, investigative reports, timelines, case summaries and partner briefings for audiences who were not present for the investigation.",
                "Developed an investigator training folder with preservation guidance, legal-process templates and service-provider references. Acquired and configured the unit's initial Cellebrite UFED.",
                "Served as a Field Training Officer from 2004 to 2023, teaching judgment, interviewing, report writing, evidence handling, policy application and problem solving in live operating conditions.",
            ],
        ),
        (
            "Real Estate Consultant",
            "eXp Realty / KW Select, South Metro MN",
            "June 2024 - March 2026",
            [
                "Managed client education, documents, deadlines, negotiations and issue resolution while supporting $3.2 million in 2025 residential sales.",
            ],
        ),
    ]
    for title, employer, dates, bullets in roles:
        add_job(doc, title, employer, dates)
        for i, bullet in enumerate(bullets):
            add_bullet(doc, bullet, keep_next=i == 0)

    add_section(doc, "Education")
    add_degree(
        doc,
        "Master of Arts, Police Leadership, Administration and Education",
        "University of St. Thomas, St. Paul, MN",
        "GPA: 3.94 | 2005",
    )
    add_degree(
        doc,
        "Bachelor of Arts, Criminal Justice, Magna Cum Laude",
        "St. Cloud State University, St. Cloud, MN",
        "GPA: 3.51 | 1998",
    )
    add_degree(
        doc,
        "Associate of Arts, Criminal Justice, Magna Cum Laude",
        "St. Cloud State University, St. Cloud, MN",
        "GPA: 3.50 | 1996",
    )

    add_section(doc, "Credentials, Technology and Military Service")
    add_body(
        doc,
        "Digital forensics: Certified Cyber Crime Investigator, NW3C, 2023 | Cellebrite CCO/CPA, recertified 2020 | FBI CAST cell-site analysis, 2017 | FTK, 2017 | X-Ways, 2018. Tools: Cellebrite UFED, Magnet AXIOM, FTK, X-Ways Forensics, GrayKey, LexisNexis Accurint and ZetX TRAXi. U.S. Army veteran; Reserve, Active Duty and Minnesota Army National Guard; honorably discharged.",
    )
    doc.save(path)
    sanitize_ooxml_font_defaults(path)


def add_cover_line(doc, text, size=12, bold=False, after=0, keep_next=False):
    p = doc.add_paragraph()
    set_para(p, after=after, line=1.0, keep_next=keep_next)
    style_run(p.add_run(text), size=size, bold=bold)
    return p


def build_cover(path):
    doc = new_branded_document(12)
    props = doc.core_properties
    props.title = "Troy Hokanson Cover Letter, Hugging Face Wild Card"
    props.subject = "Hugging Face Wild Card"
    props.author = TROY_NAME

    p = doc.add_paragraph()
    set_para(p, before=12, after=6, line=1.0, keep_next=True)
    style_run(p.add_run("August 8, 2026"), size=12)

    add_cover_line(doc, "Hugging Face Hiring Team", bold=True, keep_next=True)
    add_cover_line(doc, "Wild Card", after=6, keep_next=True)
    add_cover_line(doc, "Dear Hugging Face Hiring Team:", after=6, keep_next=True)

    paragraphs = [
        "The Wild Card posting fits because Hugging Face does not currently have a title that cleanly describes the work I could do. I am not a machine-learning engineer. My value sits where powerful technology meets investigators, public agencies, educators and people who need complicated systems to make sense in the real world. During 25 years in Minnesota law enforcement, I connected device artifacts, cloud records, telecommunications data, financial records, interviews and legal-process returns. Within a 10-agency task-force structure, I processed 5,304 GB of digital evidence in 2020 and wrote findings for prosecutors, agency leaders and investigators.",
        "I applied that evidence-first approach to a source-controlled career-search and application system. A private pipeline checks 16 sources using explicit location, travel, compensation, fit, privacy and duplicate-suppression rules. I review changes for false positives, missed opportunities and unsupported conclusions. The point is an explainable system that is less likely to quietly make the wrong decision.",
        "I taught Criminal Justice remotely for 18 years and served as a Field Training Officer for 19 years. Both roles required practical instruction, written feedback and repeatable decision frameworks. I could apply that perspective to misuse investigations, public-sector enablement, investigator education, operational documentation or community feedback.",
        "There are honest gaps. I have not maintained an open-source ML library, trained models at scale or worked inside a platform trust-and-safety team. I have investigated technology-enabled wrongdoing, examined digital evidence, tested AI-assisted workflows and explained difficult material without hiding behind jargon. That would add disciplined operational and user perspective without replacing engineering judgment. If Hugging Face sees a need for that combination, I would welcome a practical conversation.",
    ]
    for text in paragraphs:
        p = doc.add_paragraph()
        set_para(p, after=6, line=1.05)
        style_run(p.add_run(text), size=12)

    p = doc.add_paragraph()
    set_para(p, before=3, after=12, line=1.0, keep_next=True)
    style_run(p.add_run("Respectfully,"), size=12)
    add_cover_line(doc, TROY_NAME, size=12)

    doc.save(path)
    sanitize_ooxml_font_defaults(path)


if __name__ == "__main__":
    OUT.mkdir(parents=True, exist_ok=True)
    build_resume(OUT / "Troy_Hokanson_Resume_Hugging_Face_Wild_Card.docx")
    build_cover(OUT / "Troy_Hokanson_Cover_Letter_Hugging_Face_Wild_Card.docx")
    print(OUT)
