#!/usr/bin/env python3

from __future__ import annotations

import json
import re
import sys
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

REPO_ROOT = Path(__file__).resolve().parents[2]
sys.path.insert(0, str(REPO_ROOT))

import docx_header
from docx_header import add_paragraph_bottom_border, build_navy_header


APP_DIR = Path(__file__).resolve().parent
OUTPUT_DIR = APP_DIR / "output"
BUILD_LOG_DIR = APP_DIR / "build_logs"
FONT = "EB Garamond"
BLACK = RGBColor(0x14, 0x14, 0x14)
GRAY = RGBColor(0x55, 0x55, 0x55)
STEEL = RGBColor(0x2D, 0x6A, 0x9F)


RESUME = {
    "summary": (
        "Investigator and OSINT analyst transitioning to corporate personnel-risk work after 25 years of sworn "
        "service, including 5.5 years in detective and digital-forensics assignments. Managed complex investigations "
        "from intake through reporting, corroborated people and organizations across public records, LexisNexis "
        "Accurint, interviews, financial and provider records, and digital evidence, and prepared findings for legal "
        "and command decisions. Trained in background investigation methods and trusted with sensitive data."
    ),
    "capabilities": (
        "OSINT and public records | LexisNexis Accurint | Background research and identity resolution | Source "
        "corroboration | Intake through reporting | Personnel-risk assessment | Sensitive data and privacy | Analytical "
        "reporting | Research playbooks | Digital evidence and timelines | Legal and executive briefings"
    ),
    "page1_jobs": [
        {
            "title": "Real Estate Consultant",
            "dates": "June 2024 - March 2026",
            "employer": "Residential Real Estate | South Metro Minnesota",
            "bullets": [
                "Protected confidential client and financial information while researching property records and coordinating time-sensitive work across clients and industry partners. Completed $3.2 million in residential sales."
            ],
        },
        {
            "title": "Police Officer",
            "dates": "January 2022 - May 2024",
            "employer": "Lakeville Police Department | Lakeville, Minnesota",
            "bullets": [
                "Handled conflicts, threats, harassment, and safety matters requiring interviews, behavioral assessment, accurate documentation, discretion, and timely escalation while mentoring officers on evidence and policy."
            ],
        },
        {
            "title": "Detective / Digital Forensic Examiner",
            "dates": "June 2017 - December 2021",
            "employer": "Dakota County Electronic Crimes Task Force, assigned from Lakeville Police Department | Minnesota",
            "bullets": [
                "Served as Lakeville Police Department's digital-forensics resource in a ten-agency task force, independently managing complex investigations with investigators, technical personnel, and attorneys.",
                "Corroborated leads across social media, public and business records, LexisNexis Accurint, interviews, financial and provider records, and digital evidence, documenting uncertainty.",
                "Processed 5,304 GB of digital evidence in 2020 using Cellebrite, GrayKey, X-Ways Forensics, and Magnet AXIOM. Converted communications and account records into findings for non-technical leaders.",
            "Performed write-blocked acquisition and FTK/IEF analysis of a computer in a wire-fraud matter, identified responsive artifacts, and produced report media for the police case and the Minnesota Commerce Fraud Bureau."
            ],
        },
        {
            "title": "Detective / Electronic Crimes Unit",
            "dates": "September 2016 - June 2017",
            "employer": "Lakeville Police Department | Lakeville, Minnesota",
            "bullets": [
                "Built an investigator resource with preservation-request, administrative-subpoena, and search-warrant templates plus service-provider guidance, then helped investigators apply consistent research and evidence standards without overstating what the information established."
            ],
        },
    ],
    "page2_jobs": [
        {
            "title": "Police Officer / Investigator",
            "dates": "March 2010 - May 2011",
            "employer": "Lakeville Police Department | Lakeville, Minnesota",
            "bullets": [
                "Investigated employee-related occupational fraud involving approximately $80,000 in unauthorized company-card charges. Organized receipt and transaction evidence in Excel and prepared a documented case package that supported a felony conviction and court-ordered restitution."
            ],
        },
        {
            "title": "Police Officer / Field Training Officer",
            "dates": "June 2011 - August 2016",
            "employer": "Lakeville Police Department | Lakeville, Minnesota",
            "bullets": [
                "Co-founded a probation-liaison program with county community-corrections personnel, combining records, field observations, human reporting, and compliance information to prioritize joint activity and document results."
            ],
        },
        {
            "title": "Police Officer / Field Training Officer",
            "dates": "November 1998 - February 2010",
            "employer": "Lakeville Police Department | Lakeville, Minnesota",
            "bullets": [
                "Served as a Field Training Officer across 18 years of the career, training and evaluating officers in interviewing, report writing, policy, evidence handling, technology, privacy, and sound decisions under pressure.",
                "Coordinated an approximately $40,000 Target-funded Genetec AutoVu project with department leadership, city IT, BCA CJIS, and Genetec, documenting the use case, access controls, and operating workflow."
            ],
        },
        {
            "title": "Adjunct Faculty / Criminal Justice",
            "dates": "March 2007 - October 2025",
            "employer": "University of Phoenix | Remote, concurrent with sworn service",
            "bullets": [
                "Taught undergraduate Criminal Justice courses remotely for 18 years, evaluated source-supported written work, and explained investigative, legal, ethics, and evidence concepts to learners with varied experience."
            ],
        },
        {
            "title": "U.S. Army",
            "dates": "8 years 3 months",
            "employer": "Reserve, Active Duty, and Minnesota Army National Guard | Honorably Discharged",
            "bullets": [
                "Served in Infantry, Armor, Motor Transport, and Military Police roles. Held a Secret security clearance during active-duty service."
            ],
        },
    ],
}


COVER_PARAGRAPHS = [
    (
        "SandboxAQ's Background Investigator / Analyst role sits where I want to continue working: careful OSINT, "
        "source verification, personnel risk, and concise reporting for leaders who have to make difficult decisions. "
        "SandboxAQ works across quantitative AI, cybersecurity, navigation, life sciences, and other high-stakes fields. "
        "In that setting, personnel-risk research has to protect people, data, and intellectual property without outrunning the evidence."
    ),
    (
        "Across 25 years with the Lakeville Police Department, I carried investigations from intake through interviews, "
        "records research, digital-evidence review, corroboration, reporting, and legal handoff. As Lakeville's representative "
        "and digital-forensics resource in a ten-agency task force, I worked independently while coordinating with investigators, "
        "supervisors, technical personnel, attorneys, and outside agencies. The work required discretion, source evaluation, "
        "accurate documentation, and prompt escalation when facts were incomplete or risk was increasing."
    ),
    (
        "My research methods align closely with the role. I developed leads through public social media, public and business "
        "records, LexisNexis Accurint, interviews, financial information, provider records, device evidence, and call-detail/location "
        "data. In one investigation, I documented public profiles, mapped an apparent associate through open social content, "
        "corroborated identity information with Accurint, and initiated preservation before the information could change. In a "
        "separate Business Email Compromise investigation, I used public records, shell-company analysis, financial tracing, and "
        "cross-case pattern recognition to connect multiple victims. The matter closed with a felony conviction and more than "
        "$360,000 in verified losses."
    ),
    (
        "I also understand the value of a consistent method. In Lakeville's Electronic Crimes Unit, I built an investigator "
        "resource containing preservation-request, administrative-subpoena, and search-warrant templates with service-provider "
        "guidance. I spent 18 years as a Field Training Officer and 18 years teaching Criminal Justice remotely. Those roles "
        "required me to explain standards, review written work, identify unsupported conclusions, and help others improve without "
        "taking over the assignment."
    ),
    (
        "My background is related investigative work rather than in-house employment screening. I have formal background-investigation "
        "training and extensive experience researching U.S.-based individuals and organizations, but I have not administered FCRA "
        "screening or worked inside a corporate facility-clearance program. That is the real gap. What transfers directly is the "
        "ability to verify accuracy, assess source reliability, handle adverse information fairly, document a defensible risk "
        "determination, and separate established facts from inference. I am also comfortable learning a repeatable scoring method "
        "and working within the privacy, Legal, Human Resources, and Information Security controls SandboxAQ sets."
    ),
    (
        "Following retirement from sworn service, I confirmed that investigations, analytical writing, and research-method "
        "development are where I do my best work. I would welcome a candid discussion about how that experience could support "
        "SandboxAQ's insider threat and personnel-risk program as it matures."
    ),
]


def set_run_font(run, size: float, *, bold: bool = False, italic: bool = False, color=BLACK) -> None:
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        r_fonts.set(qn(f"w:{attr}"), FONT)


def set_paragraph(paragraph, *, before=0, after=0, line=1.05, keep_next=False, keep_together=False) -> None:
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    pf.widow_control = True
    pf.keep_with_next = keep_next
    pf.keep_together = keep_together


def prepare_document(doc_type: str) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    if doc_type == "resume":
        left = right = 0.65
        bottom = 0.55
        top = 1.46
        body_size = 10.25
    else:
        left = right = 0.78
        bottom = 0.68
        top = 1.52
        body_size = 10.25
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(body_size)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.05
    list_style = doc.styles["List Bullet"]
    list_style.font.name = FONT
    list_style.font.size = Pt(body_size)
    docx_header.NAME_FONT = FONT
    docx_header.CONTACT_FONT = FONT
    docx_header.BODY_FONT = FONT
    docx_header.CONTACT_PARTS = [
        item for item in docx_header.CONTACT_PARTS if item[0] != docx_header.TROY_LOCATION
    ]
    build_navy_header(
        doc,
        body_top_margin_inches=top,
        body_bottom_margin_inches=bottom,
        body_left_margin_inches=left,
        body_right_margin_inches=right,
    )
    section.different_first_page_header_footer = False
    replace_header_with_table(section.header)
    return doc


def replace_header_with_table(header) -> None:
    """Build a renderer-stable, full-width navy header with normal table flow."""
    for child in list(header._element):
        header._element.remove(child)
    table = header.add_table(rows=1, cols=1, width=Inches(8.5))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    tbl_w.set(qn("w:w"), "12240")
    tbl_w.set(qn("w:type"), "dxa")
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)
    row = table.rows[0]
    row.height = Inches(1.28)
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    cell = row.cells[0]
    cell.width = Inches(8.5)
    docx_header.shade_cell(cell, "0D1B2A")
    docx_header.remove_cell_borders(cell)
    docx_header.set_cell_margins(cell, top=110, bottom=80, left=0, right=0)

    name_p = cell.paragraphs[0]
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_p.paragraph_format.left_indent = Inches(1.91)
    name_p.paragraph_format.right_indent = Inches(1.91)
    docx_header.set_paragraph_format(name_p, before=7, after=4, line=1.0)
    add_paragraph_bottom_border(name_p, color_hex="C9A84C", size=7, space=3)
    docx_header.set_run(
        name_p.add_run("Troy Hokanson"),
        font=FONT,
        size=26,
        bold=True,
        color=RGBColor(0xFF, 0xFF, 0xFF),
    )

    contact_p = cell.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    docx_header.set_paragraph_format(contact_p, before=3, after=0, line=1.0)
    for index, (label, url) in enumerate(docx_header.CONTACT_PARTS):
        if index:
            docx_header.set_run(
                contact_p.add_run(" | "),
                font=FONT,
                size=10,
                color=RGBColor(0xC9, 0xA8, 0x4C),
            )
        if url:
            docx_header.add_hyperlink(
                contact_p,
                label,
                url,
                color=RGBColor(0xC9, 0xA8, 0x4C),
                font=FONT,
                size=10,
            )
        else:
            docx_header.set_run(
                contact_p.add_run(label),
                font=FONT,
                size=10,
                color=RGBColor(0xC9, 0xA8, 0x4C),
            )


def add_section_heading(doc: Document, text: str, *, page_break_before: bool = False) -> None:
    p = doc.add_paragraph()
    set_paragraph(p, before=10, after=4, line=1.0, keep_next=True)
    p.paragraph_format.page_break_before = page_break_before
    set_run_font(p.add_run(text.upper()), 11, bold=True, color=STEEL)
    add_paragraph_bottom_border(p, color_hex="C9A84C", size=6)


def add_body_paragraph(doc: Document, text: str, *, size=10.25, after=4) -> None:
    p = doc.add_paragraph()
    set_paragraph(p, before=0, after=after, line=1.05, keep_together=True)
    set_run_font(p.add_run(text), size)


def add_job(doc: Document, job: dict) -> None:
    p = doc.add_paragraph()
    set_paragraph(p, before=6, after=1, line=1.0, keep_next=True, keep_together=True)
    set_run_font(p.add_run(job["title"]), 10.5, bold=True)
    set_run_font(p.add_run(" | " + job["dates"]), 10.0, bold=True, color=GRAY)
    p = doc.add_paragraph()
    set_paragraph(p, before=0, after=2, line=1.0, keep_next=True, keep_together=True)
    set_run_font(p.add_run(job["employer"]), 9.75, italic=True, color=GRAY)
    for index, bullet in enumerate(job["bullets"]):
        p = doc.add_paragraph(style="List Bullet")
        set_paragraph(p, before=0, after=1, line=1.03, keep_together=True, keep_next=False)
        p.paragraph_format.left_indent = Inches(0.22)
        p.paragraph_format.first_line_indent = Inches(-0.14)
        for run in p.runs:
            run.text = ""
        set_run_font(p.add_run(bullet), 10.25)


def add_degree(doc: Document, lines: list[str]) -> None:
    p = doc.add_paragraph()
    set_paragraph(p, before=3, after=1, line=1.0, keep_together=True)
    for index, line in enumerate(lines):
        run = p.add_run(line)
        set_run_font(run, 9.75, bold=index == 0)
        if index < len(lines) - 1:
            run.add_break()


def build_resume() -> Path:
    doc = prepare_document("resume")
    add_section_heading(doc, "Professional Summary")
    add_body_paragraph(doc, RESUME["summary"], after=4)
    add_section_heading(doc, "Background Investigation and Analytical Capabilities")
    add_body_paragraph(doc, RESUME["capabilities"], after=4)
    add_section_heading(doc, "Professional Experience")
    for job in RESUME["page1_jobs"]:
        add_job(doc, job)
    add_section_heading(doc, "Additional Experience")
    for job in RESUME["page2_jobs"]:
        add_job(doc, job)
    add_section_heading(doc, "Education")
    add_degree(doc, [
        "Master of Arts, Police Leadership, Administration and Education",
        "University of St. Thomas, St. Paul, MN | GPA: 3.94 | 2005",
    ])
    add_degree(doc, [
        "Bachelor of Arts, Criminal Justice, Magna Cum Laude",
        "St. Cloud State University, St. Cloud, MN | GPA: 3.51 | 1998",
    ])
    add_degree(doc, [
        "Associate of Arts, Criminal Justice, Magna Cum Laude",
        "St. Cloud State University, St. Cloud, MN | GPA: 3.50 | 1996",
    ])
    add_section_heading(doc, "Selected Training and Credentials")
    add_body_paragraph(
        doc,
        "Background Investigation, St. Paul Police Department, May 2006 | Certified Fraud Examiner, actively pursuing through ACFE, 2026 | Certified Cyber Crime Investigator, 2023 | Reid Technique of Interviewing and Interrogation",
        size=9.75,
        after=0,
    )
    output = OUTPUT_DIR / "Hokanson_Resume_SandboxAQ_Background_Investigator_Analyst.docx"
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    sanitize_docx_fonts(output)
    return output


def add_cover_line(doc: Document, text: str, *, after=0, bold=False) -> None:
    p = doc.add_paragraph()
    set_paragraph(p, before=0, after=after, line=1.0, keep_together=True)
    set_run_font(p.add_run(text), 10.25, bold=bold)


def build_cover() -> Path:
    doc = prepare_document("cover")
    add_cover_line(doc, "August 21, 2026", after=4)
    add_cover_line(doc, "SandboxAQ Hiring Team", after=4)
    add_cover_line(doc, "Dear SandboxAQ Hiring Team,", after=5)
    for paragraph in COVER_PARAGRAPHS:
        p = doc.add_paragraph()
        set_paragraph(p, before=0, after=3, line=1.0, keep_together=False)
        set_run_font(p.add_run(paragraph), 10.25)
    p = doc.add_paragraph()
    set_paragraph(p, before=5, after=0, line=1.0, keep_next=True)
    set_run_font(p.add_run("Respectfully,"), 10.25)
    p = doc.add_paragraph()
    set_paragraph(p, before=8, after=0, line=1.0, keep_together=True)
    set_run_font(p.add_run("Troy Hokanson"), 10.25)
    output = OUTPUT_DIR / "Hokanson_Cover_SandboxAQ_Background_Investigator_Analyst.docx"
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    sanitize_docx_fonts(output)
    return output


def sanitize_docx_fonts(path: Path) -> None:
    replacements = {
        "Calibri Light": FONT,
        "Calibri": FONT,
        "Aptos Display": FONT,
        "Aptos": FONT,
        "Arial": FONT,
    }
    temporary = path.with_suffix(".fonts.docx")
    with zipfile.ZipFile(path, "r") as source, zipfile.ZipFile(temporary, "w", zipfile.ZIP_DEFLATED) as target:
        for item in source.infolist():
            data = source.read(item.filename)
            if item.filename.startswith("word/") and item.filename.endswith((".xml", ".rels")):
                text = data.decode("utf-8", errors="replace")
                for old, new in replacements.items():
                    text = re.sub(re.escape(old), new, text, flags=re.IGNORECASE)
                data = text.encode("utf-8")
            target.writestr(item, data)
    temporary.replace(path)


def inspect_docx(path: Path, doc_type: str) -> dict:
    doc = Document(path)
    section = doc.sections[0]
    expected = {
        "resume": {"left": 0.65, "right": 0.65, "bottom": 0.55, "top": 1.46},
        "cover": {"left": 0.78, "right": 0.78, "bottom": 0.68, "top": 1.52},
    }[doc_type]
    bullets = [p for p in doc.paragraphs if p.style and p.style.name == "List Bullet"]
    keep_next = sum(1 for p in doc.paragraphs if p.paragraph_format.keep_with_next)
    with zipfile.ZipFile(path) as archive:
        xml = "\n".join(
            archive.read(name).decode("utf-8", errors="replace")
            for name in archive.namelist()
            if name.startswith("word/") and name.endswith(".xml")
        ).lower()
        rels = "\n".join(
            archive.read(name).decode("utf-8", errors="replace")
            for name in archive.namelist()
            if name.startswith("word/") and name.endswith(".rels")
        )
    actual = {
        "left": round(section.left_margin.inches, 3),
        "right": round(section.right_margin.inches, 3),
        "bottom": round(section.bottom_margin.inches, 3),
        "top": round(section.top_margin.inches, 3),
    }
    checks = {
        "us_letter": round(section.page_width.inches, 2) == 8.5 and round(section.page_height.inches, 2) == 11,
        "margins_match": actual == expected,
        "header_part_present": bool(doc.sections[0].header.paragraphs or doc.sections[0].header.tables),
        "real_list_style_present": len(bullets) > 0 if doc_type == "resume" else True,
        "keep_with_next_present": keep_next > 0,
        "garamond_present": "eb garamond" in xml,
        "calibri_absent": "calibri" not in xml,
        "aptos_absent": "aptos" not in xml,
        "arial_absent": "arial" not in xml,
        "contact_hyperlinks_present": "hyperlink" in rels.lower(),
    }
    return {
        "path": str(path),
        "document_type": doc_type,
        "section_count": len(doc.sections),
        "paragraph_count": len(doc.paragraphs),
        "bullet_paragraph_count": len(bullets),
        "keep_with_next_count": keep_next,
        "margins_inches": actual,
        "checks": checks,
        "passed": all(checks.values()),
    }


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    BUILD_LOG_DIR.mkdir(parents=True, exist_ok=True)
    resume = build_resume()
    cover = build_cover()
    report = {
        "layout_contract": str(APP_DIR / "layout_contract.json"),
        "documents": [inspect_docx(resume, "resume"), inspect_docx(cover, "cover")],
    }
    report["passed"] = all(item["passed"] for item in report["documents"])
    (BUILD_LOG_DIR / "docx_structure_audit.json").write_text(json.dumps(report, indent=2), encoding="utf-8")
    print(json.dumps(report, indent=2))
    if not report["passed"]:
        raise SystemExit(1)


if __name__ == "__main__":
    main()
