#!/usr/bin/env python3

from __future__ import annotations

import json
import re
import sys
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

REPO_ROOT = Path(__file__).resolve().parents[2]
sys.path.insert(0, str(REPO_ROOT))

import docx_header
from docx_header import add_paragraph_bottom_border, build_navy_header


APP_DIR = Path(__file__).resolve().parent
OUTPUT_DIR = APP_DIR / "output"
BUILD_LOG_DIR = APP_DIR / "build_logs"
FONT = "EB Garamond"
BLACK = RGBColor(0x14, 0x14, 0x14)
GRAY = RGBColor(0x55, 0x55, 0x55)
STEEL = RGBColor(0x2D, 0x6A, 0x9F)


RESUME = {
    "summary": (
        "Customer-facing technology and operations professional with 25 years of public-safety "
        "service, 18 years of remote college instruction, and recent private-sector client management. "
        "Led an agency-side $40,000 Genetec AutoVu ALPR implementation, coordinated technical priorities "
        "across a ten-agency task force, and helped daily users adopt fleet video, mobile forensics, "
        "communications, and digital-evidence workflows. Translates operational needs into practical "
        "training, documented processes, and clear decisions for executives, technical partners, and "
        "frontline users."
    ),
    "capabilities": (
        "Customer adoption and enablement | Operational discovery and workshops | Workflow assessment "
        "and change | Technical implementation support | Stakeholder alignment | Executive and frontline "
        "communication | Vendor, IT, and agency coordination | Issue resolution and escalation | Training "
        "and mentoring | Fleet video, ALPR, and digital evidence | Microsoft 365"
    ),
    "page1_jobs": [
        {
            "title": "Real Estate Consultant",
            "dates": "June 2024 - March 2026",
            "employer": "Residential Real Estate | South Metro Minnesota",
            "bullets": [
                "Managed client relationships from initial consultation through negotiation, inspection, financing, title, and closing; completed $3.2M in residential sales during the transition from law enforcement.",
                "Coordinated clients, lenders, inspectors, appraisers, title professionals, and other agents to identify barriers, explain options, and keep time-sensitive transactions moving.",
            ],
        },
        {
            "title": "Police Officer",
            "dates": "January 2022 - May 2024",
            "employer": "Lakeville Police Department | Lakeville, Minnesota",
            "bullets": [
                "Used synchronized Axon Body 3 and Fleet 2 systems, Motorola radios, Microsoft 365, and related tools in daily patrol and evidence workflows.",
                "Helped officers and supervisors resolve policy, evidence, and technology questions while maintaining safety, documentation, and public-service standards.",
            ],
        },
        {
            "title": "Detective / Digital Forensic Examiner",
            "dates": "September 2016 - December 2021",
            "employer": "Lakeville Police Department and Dakota County Electronic Crimes Task Force | Minnesota",
            "bullets": [
                "Assigned to the ten-agency task force from June 2017 through December 2021, coordinating priorities, technical guidance, and examination support across partner agencies.",
                "Processed 5,304 GB of digital evidence in 2020 using Cellebrite, GrayKey, Magnet AXIOM, X-Ways, and related platforms; converted technical findings into reports and briefings for investigators, supervisors, and legal decision-makers.",
                "Used the unit's existing Cellebrite UFED, built repeatable training resources, and supported daily users through complex device and evidence questions.",
            ],
        },
    ],
    "page2_jobs": [
        {
            "title": "Police Officer / Field Training Officer",
            "dates": "November 1998 - August 2016",
            "employer": "Lakeville Police Department | Lakeville, Minnesota",
            "bullets": [
                "Served as agency-side operational lead for a roughly $40,000 Target-funded Genetec AutoVu ALPR project, coordinating the use case with Target, department leadership, city IT, BCA CJIS, and Genetec.",
                "Helped establish the secure garage Wi-Fi workflow for nightly records updates and worked with Genetec to troubleshoot camera and connectivity issues during implementation and use.",
                "Served 19 years as a Field Training Officer and held supplemental assignments in reserve officer development, park ranger hiring and training, and juvenile traffic-diversion instruction.",
            ],
        },
        {
            "title": "Adjunct Faculty / Criminal Justice",
            "dates": "March 2007 - October 2025",
            "employer": "University of Phoenix | Remote, concurrent with sworn service",
            "bullets": [
                "Taught undergraduate Criminal Justice courses remotely for 18 years, converting complex legal, investigative, and technical subjects into structured lessons for adult learners.",
                "Received the Phoenix500 Faculty Excellence Award in 2020 and 2021 and a Faculty of the Year nomination in 2021.",
            ],
        },
        {
            "title": "U.S. Army",
            "dates": "8 years 3 months",
            "employer": "Reserve, Active Duty, and Minnesota Army National Guard | Honorably Discharged",
            "bullets": [
                "Served in Infantry, Armor, Motor Transport, and Military Police roles with direct responsibility for equipment, safety, team coordination, and mission execution.",
            ],
        },
    ],
}


COVER_PARAGRAPHS = [
    (
        "Samsara's Strategic Customer Success Manager role is about understanding physical "
        "operations, removing barriers, and helping large fleet customers get measurable value from "
        "safety technology. That is why I am applying. I have spent much of my career on the customer "
        "side of connected systems, where a product succeeds only when it works in the field and users "
        "trust the workflow."
    ),
    (
        "In 2007, I served as agency-side operational lead for a roughly $40,000 Target-funded Genetec "
        "AutoVu ALPR project. I coordinated Target, department leadership, city IT, BCA CJIS, and Genetec, "
        "helped establish secure Wi-Fi for nightly records updates, and worked with Genetec on camera and "
        "connectivity problems. That was hands-on discovery, implementation, adoption, vendor coordination, "
        "and field troubleshooting."
    ),
    (
        "Later, as Lakeville's representative on a ten-agency electronic crimes task force, I coordinated "
        "technical priorities, personally processed 5,304 GB of digital evidence in 2020, and translated "
        "findings for supervisors and legal decision-makers. Nineteen years as a Field Training Officer and "
        "18 years teaching remotely taught me to assess what users understand, adjust the explanation, "
        "document the process, and follow through."
    ),
    (
        "Real estate added private-sector client management. I completed $3.2M in sales while coordinating "
        "clients and transaction partners through time-sensitive decisions. The work required honest "
        "conversations, careful priority management, and steady communication when a problem threatened "
        "the timeline."
    ),
    (
        "I have not owned an Enterprise SaaS portfolio or formal renewal targets, and I will not dress up "
        "adjacent work as direct CSM tenure. The fit rests on customer-side implementation, technical "
        "credibility, executive communication, training, and long-tested judgment when operations and "
        "technology do not line up cleanly. Samsara's Video-Based Safety, Vehicle Telematics, and Connected "
        "Workflows products make this a transition worth pursuing."
    ),
]


def set_run_font(run, size: float, *, bold: bool = False, italic: bool = False, color=BLACK) -> None:
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        r_fonts.set(qn(f"w:{attr}"), FONT)


def set_paragraph(paragraph, *, before=0, after=0, line=1.05, keep_next=False, keep_together=False) -> None:
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    pf.widow_control = True
    pf.keep_with_next = keep_next
    pf.keep_together = keep_together


def prepare_document(doc_type: str) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)

    if doc_type == "resume":
        left = right = 0.65
        bottom = 0.55
        top = 1.46
        body_size = 10.25
    else:
        left = right = 0.78
        bottom = 0.68
        top = 1.52
        body_size = 10.5

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(body_size)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.05

    list_style = doc.styles["List Bullet"]
    list_style.font.name = FONT
    list_style.font.size = Pt(body_size)

    docx_header.NAME_FONT = FONT
    docx_header.CONTACT_FONT = FONT
    docx_header.BODY_FONT = FONT
    docx_header.CONTACT_PARTS = [
        item for item in docx_header.CONTACT_PARTS if item[0] != docx_header.TROY_LOCATION
    ]
    build_navy_header(
        doc,
        body_top_margin_inches=top,
        body_bottom_margin_inches=bottom,
        body_left_margin_inches=left,
        body_right_margin_inches=right,
    )
    return doc


def add_section_heading(doc: Document, text: str, *, page_break_before: bool = False) -> None:
    p = doc.add_paragraph()
    set_paragraph(p, before=14, after=6, line=1.0, keep_next=True)
    p.paragraph_format.page_break_before = page_break_before
    run = p.add_run(text.upper())
    set_run_font(run, 11, bold=True, color=STEEL)
    add_paragraph_bottom_border(p, color_hex="C9A84C", size=6)


def add_body_paragraph(doc: Document, text: str, *, size=10.25, after=4) -> None:
    p = doc.add_paragraph()
    set_paragraph(p, before=0, after=after, line=1.05, keep_together=True)
    set_run_font(p.add_run(text), size)


def add_job(doc: Document, job: dict) -> None:
    p = doc.add_paragraph()
    set_paragraph(p, before=8, after=2, line=1.0, keep_next=True, keep_together=True)
    set_run_font(p.add_run(job["title"]), 10.5, bold=True)
    set_run_font(p.add_run(" | " + job["dates"]), 10.0, bold=True, color=GRAY)

    p = doc.add_paragraph()
    set_paragraph(p, before=0, after=4, line=1.0, keep_next=True, keep_together=True)
    set_run_font(p.add_run(job["employer"]), 9.75, italic=True, color=GRAY)

    for index, bullet in enumerate(job["bullets"]):
        p = doc.add_paragraph(style="List Bullet")
        set_paragraph(
            p,
            before=0,
            after=2,
            line=1.05,
            keep_together=True,
            keep_next=index < len(job["bullets"]) - 1,
        )
        p.paragraph_format.left_indent = Inches(0.22)
        p.paragraph_format.first_line_indent = Inches(-0.14)
        for run in p.runs:
            run.text = ""
        set_run_font(p.add_run(bullet), 10.25)


def add_degree(doc: Document, lines: list[str]) -> None:
    p = doc.add_paragraph()
    set_paragraph(p, before=5, after=2, line=1.0, keep_together=True)
    for index, line in enumerate(lines):
        run = p.add_run(line)
        set_run_font(run, 9.75, bold=index == 0)
        if index < len(lines) - 1:
            run.add_break()


def build_resume() -> Path:
    doc = prepare_document("resume")
    add_section_heading(doc, "Professional Summary")
    add_body_paragraph(doc, RESUME["summary"], after=4)
    add_section_heading(doc, "Customer and Technical Capabilities")
    add_body_paragraph(doc, RESUME["capabilities"], after=4)
    add_section_heading(doc, "Professional Experience")
    for job in RESUME["page1_jobs"]:
        add_job(doc, job)

    add_section_heading(doc, "Additional Experience", page_break_before=True)
    for job in RESUME["page2_jobs"]:
        add_job(doc, job)

    add_section_heading(doc, "Education")
    add_degree(doc, [
        "Master of Arts, Police Leadership, Administration and Education",
        "University of St. Thomas, St. Paul, MN",
        "GPA: 3.94",
        "2005",
    ])
    add_degree(doc, [
        "Bachelor of Arts, Criminal Justice, Magna Cum Laude",
        "St. Cloud State University, St. Cloud, MN",
        "GPA: 3.51",
        "1998",
    ])
    add_degree(doc, [
        "Associate of Arts, Criminal Justice, Magna Cum Laude",
        "St. Cloud State University, St. Cloud, MN",
        "GPA: 3.50",
        "1996",
    ])

    add_section_heading(doc, "Selected Training and Credentials")
    add_body_paragraph(
        doc,
        "University of Phoenix Certified Advanced Facilitator | BCA Law Enforcement Supervision "
        "and Management, 98 hours | Cellebrite Certified Operator and Physical Analyst training | "
        "NW3C cybercrime investigation training",
        size=9.75,
        after=0,
    )

    output = OUTPUT_DIR / "Hokanson_Resume_Samsara_Strategic_Customer_Success_Manager.docx"
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    sanitize_docx_fonts(output)
    return output


def add_cover_line(doc: Document, text: str, *, after=0, bold=False) -> None:
    p = doc.add_paragraph()
    set_paragraph(p, before=0, after=after, line=1.0, keep_together=True)
    set_run_font(p.add_run(text), 10.5, bold=bold)


def build_cover() -> Path:
    doc = prepare_document("cover")
    add_cover_line(doc, "August 13, 2026", after=8)
    add_cover_line(doc, "Hiring Manager")
    add_cover_line(doc, "Samsara, Inc.", after=9)
    add_cover_line(doc, "Dear Hiring Manager,", after=9)

    for paragraph in COVER_PARAGRAPHS:
        p = doc.add_paragraph()
        set_paragraph(p, before=0, after=8, line=1.05, keep_together=True)
        set_run_font(p.add_run(paragraph), 10.5)

    p = doc.add_paragraph()
    set_paragraph(p, before=5, after=0, line=1.0, keep_next=True)
    set_run_font(p.add_run("Respectfully,"), 10.5)
    p = doc.add_paragraph()
    set_paragraph(p, before=38, after=0, line=1.0, keep_together=True)
    set_run_font(p.add_run("Troy Hokanson"), 10.5)

    output = OUTPUT_DIR / "Hokanson_Cover_Samsara_Strategic_Customer_Success_Manager.docx"
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    sanitize_docx_fonts(output)
    return output


def sanitize_docx_fonts(path: Path) -> None:
    replacements = {
        "Calibri Light": FONT,
        "Calibri": FONT,
        "Aptos Display": FONT,
        "Aptos": FONT,
        "Arial": FONT,
    }
    temporary = path.with_suffix(".fonts.docx")
    with zipfile.ZipFile(path, "r") as source, zipfile.ZipFile(temporary, "w", zipfile.ZIP_DEFLATED) as target:
        for item in source.infolist():
            data = source.read(item.filename)
            if item.filename.startswith("word/") and item.filename.endswith((".xml", ".rels")):
                text = data.decode("utf-8", errors="replace")
                for old, new in replacements.items():
                    text = re.sub(re.escape(old), new, text, flags=re.IGNORECASE)
                data = text.encode("utf-8")
            target.writestr(item, data)
    temporary.replace(path)


def inspect_docx(path: Path, doc_type: str) -> dict:
    doc = Document(path)
    section = doc.sections[0]
    expected = {
        "resume": {"left": 0.65, "right": 0.65, "bottom": 0.55, "top": 1.46},
        "cover": {"left": 0.78, "right": 0.78, "bottom": 0.68, "top": 1.52},
    }[doc_type]

    bullet_paragraphs = [p for p in doc.paragraphs if p.style and p.style.name == "List Bullet"]
    keep_next = sum(1 for p in doc.paragraphs if p.paragraph_format.keep_with_next)
    with zipfile.ZipFile(path) as archive:
        xml = "\n".join(
            archive.read(name).decode("utf-8", errors="replace")
            for name in archive.namelist()
            if name.startswith("word/") and name.endswith(".xml")
        ).lower()
        rels = "\n".join(
            archive.read(name).decode("utf-8", errors="replace")
            for name in archive.namelist()
            if name.startswith("word/") and name.endswith(".rels")
        )

    actual = {
        "left": round(section.left_margin.inches, 3),
        "right": round(section.right_margin.inches, 3),
        "bottom": round(section.bottom_margin.inches, 3),
        "top": round(section.top_margin.inches, 3),
    }
    checks = {
        "us_letter": round(section.page_width.inches, 2) == 8.5 and round(section.page_height.inches, 2) == 11,
        "margins_match": actual == expected,
        "header_part_present": bool(doc.sections[0].header.paragraphs),
        "real_list_style_present": len(bullet_paragraphs) > 0 if doc_type == "resume" else True,
        "keep_with_next_present": keep_next > 0,
        "garamond_present": "eb garamond" in xml,
        "calibri_absent": "calibri" not in xml,
        "aptos_absent": "aptos" not in xml,
        "arial_absent": "arial" not in xml,
        "contact_hyperlinks_present": "hyperlink" in rels.lower(),
    }
    return {
        "path": str(path),
        "document_type": doc_type,
        "section_count": len(doc.sections),
        "paragraph_count": len(doc.paragraphs),
        "bullet_paragraph_count": len(bullet_paragraphs),
        "keep_with_next_count": keep_next,
        "margins_inches": actual,
        "checks": checks,
        "passed": all(checks.values()),
    }


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    BUILD_LOG_DIR.mkdir(parents=True, exist_ok=True)
    resume = build_resume()
    cover = build_cover()
    report = {
        "layout_contract": str(APP_DIR / "layout_contract.json"),
        "documents": [inspect_docx(resume, "resume"), inspect_docx(cover, "cover")],
    }
    report["passed"] = all(item["passed"] for item in report["documents"])
    (BUILD_LOG_DIR / "docx_structure_audit.json").write_text(
        json.dumps(report, indent=2), encoding="utf-8"
    )
    print(json.dumps(report, indent=2))
    if not report["passed"]:
        raise SystemExit(1)


if __name__ == "__main__":
    main()
