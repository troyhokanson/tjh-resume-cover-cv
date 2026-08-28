#!/usr/bin/env python3
from __future__ import annotations

import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

REPO_ROOT = Path(__file__).resolve().parents[2]
sys.path.insert(0, str(REPO_ROOT))

from docx_header import add_paragraph_bottom_border, build_navy_header
from anti_ai_scan import scan_pdf

APP_DIR = Path(__file__).resolve().parent
OUTPUT_DIR = APP_DIR / "output"
BUILD_LOG_DIR = APP_DIR / "build_logs"
FONT = "EB Garamond"
BLACK = RGBColor(0x14, 0x14, 0x14)
GRAY = RGBColor(0x55, 0x55, 0x55)
STEEL = RGBColor(0x2D, 0x6A, 0x9F)

RESUME_SUMMARY = (
    "Public-safety technology practitioner, instructor, and former investigator with 25 years of sworn service, "
    "18 years of remote college teaching, and more than five years of hands-on digital forensics work. Built and "
    "improved agency workflows around ALPR, mobile forensics, digital evidence, training, and cross-agency case "
    "support. Experienced taking loosely defined operational problems from need identification through implementation, "
    "documentation, user instruction, troubleshooting, and sustained use. Direct end-user experience includes ALPR, "
    "Axon body-worn and fleet video, Cellebrite UFED, Magnet AXIOM, GrayKey, X-Ways Forensics, and Microsoft 365. "
    "Developing Drone as First Responder domain knowledge and able to obtain FAA Part 107."
)

CAPABILITIES = (
    "Public-safety customer enablement | Technology implementation | Agency onboarding | Workflow design and adoption | "
    "Project ownership | Stakeholder communication | User instruction | Issue resolution and escalation | Product and user "
    "feedback | Multi-agency coordination | Digital evidence systems | ALPR and video workflows | Remote instruction | "
    "DFR domain development"
)

PROJECTS = [
    ("ALPR implementation, 2007-2010", "Helped develop and support an agency partnership with Target and worked with Genetec on an agency-side AutoVu implementation in 2007. Contributed operational and investigative requirements, user adoption, and sustained use over a multi-year program."),
    ("Electronic-crimes workflow build, 2016-2017", "Acquired and configured the unit's initial Cellebrite UFED, then built a structured investigator resource folder with preservation, administrative subpoena, search-warrant, and service-provider materials so investigators had a repeatable path for electronic evidence."),
    ("Time-sensitive digital evidence project", "Independently took ownership of an urgent investigation outside normal assignment flow, reported in during off-hours to reduce evidence-loss risk, preserved rapidly changing digital evidence, and carried the matter through legal process, examination, documentation, and prosecutor handoff."),
    ("High-priority encrypted-email investigation", "Served as the agency focal point for a complex digital case, coordinated legal process across international and federal partners, obtained the necessary technical records, organized device evidence, and translated the result into a usable investigative package under schedule pressure."),
    ("Current GitHub and AI workflow projects", "Maintain a GitHub-based career and application system with reusable standards, automated privacy and anti-AI validation, repeatable build scripts, and structured tracking. The work reflects the same builder pattern: identify friction, create a repeatable process, test it, and improve it."),
]

JOBS_PAGE1 = [
    ("Real Estate Consultant", "June 2024 - March 2026", "eXp Realty / KW Select | South Metro MN", [
        "Managed client relationships from initial consultation through negotiation, inspection, financing, title, and closing. Completed $3.2M in residential sales during the transition from law enforcement.",
        "Coordinated clients, lenders, inspectors, appraisers, title professionals, and cooperating agents through time-sensitive transactions, kept milestones visible, and explained difficult decisions in plain language.",
    ]),
    ("Police Officer", "January 2022 - May 2024", "Lakeville Police Department | Lakeville, MN", [
        "Returned to frontline operations after a specialized assignment, using Axon Body 3, Axon Fleet 2, Motorola radios, Microsoft 365, and related systems in daily public-safety workflows.",
        "Helped officers and supervisors work through technology, evidence, policy, and documentation questions while maintaining operational continuity.",
    ]),
    ("Detective / Digital Forensic Examiner", "June 2017 - December 2021", "Dakota County Electronic Crimes Task Force, assigned from Lakeville Police Department | Minnesota", [
        "Served as the Lakeville Police Department representative and digital forensics subject-matter resource in a ten-agency task force, coordinating examinations, priorities, and technical guidance across partner agencies.",
        "Processed 5,304 GB of digital evidence in 2020 with Cellebrite, GrayKey, Magnet AXIOM, X-Ways Forensics, and related platforms, then translated findings into reports and briefings for investigators, supervisors, and legal decision-makers.",
        "Supported daily users with device and evidence questions, selected fit-for-purpose tools, documented repeatable workflows, and found another technical path when the first approach did not answer the operational question.",
    ]),
]

JOBS_PAGE2 = [
    ("Detective / Electronic Crimes Unit", "September 2016 - June 2017", "Lakeville Police Department | Lakeville, MN", [
        "Acquired and configured the unit's initial Cellebrite UFED and helped investigators incorporate mobile-device evidence into existing case workflows.",
        "Created reusable electronic-evidence guidance and templates that reduced the learning curve for investigators new to digital legal process.",
    ]),
    ("Police Officer / Field Training Officer", "November 1998 - August 2016", "Lakeville Police Department | Lakeville, MN", [
        "Served 18 years as a Field Training Officer, coaching officers through policy, technology, documentation, communication, and decision-making in changing field conditions.",
        "Helped develop and support the 2007-2010 ALPR partnership with Target and worked with Genetec on an agency-side AutoVu implementation, translating field and investigative needs into a practical agency workflow.",
        "Built and delivered internal instruction and contributed to academy-style onboarding projects, including reserve-officer and seasonal park-ranger training programs.",
    ]),
    ("Adjunct Faculty / Criminal Justice", "March 2007 - October 2025", "University of Phoenix | Remote, concurrent with sworn service", [
        "Taught undergraduate Criminal Justice courses remotely for 18 years, turning complex legal, investigative, and technical subjects into structured lessons for adult learners.",
        "Adjusted explanations and written feedback to help students move from concept to independent application in an online environment.",
        "Received Phoenix500 Faculty Excellence Awards in 2020 and 2021 and a Faculty of the Year nomination in 2021.",
    ]),
    ("U.S. Army", "8 years 3 months", "Reserve, Active Duty, and Minnesota Army National Guard | Honorably Discharged", [
        "Served across multiple components with responsibility for equipment accountability, safety, team coordination, training, and mission execution."
    ]),
]

COVER_PARAGRAPHS = [
    "Skydio's Northwest DFR role is unusually close to the work I want to do next. I spent 25 years in Minnesota public safety using, introducing, and helping others adopt technology under real operational pressure. The part of the posting that stands out to me is not the drone itself. It is the responsibility to take a customer's use case, manage implementation, solve adoption problems, and stay accountable until the technology becomes useful in daily operations.",
    "That pattern has followed me throughout my career. From 2007 through 2010, I helped develop and support an agency ALPR partnership with Target, including agency-side work with Genetec AutoVu in 2007. The project required translating patrol and investigative needs into a workflow that officers could actually use. During a later electronic-crimes assignment, I acquired and configured our initial Cellebrite UFED and built a structured investigator resource folder with preservation, subpoena, search-warrant, and service-provider materials. The goal was not to own a new tool. The goal was to make the workflow repeatable enough that other investigators could use it correctly without starting over every time.",
    "I have also taken ownership when the work was not neatly assigned. In one time-sensitive digital investigation, I self-initiated the case and reported in outside normal hours because evidence could disappear if nobody moved quickly. I preserved the available digital evidence, built the legal-process path, coordinated the examination, documented the work, and carried the package forward for prosecution. In another high-priority encrypted-email matter, I served as the agency focal point and coordinated across international and federal partners to obtain technical records under a tight timeline. Those cases reinforced the same lesson I see in Skydio's posting: implementation succeeds when one person owns the details, keeps the stakeholders connected, and does not let schedule or technical friction become somebody else's problem.",
    "Training and adoption are equally familiar. I served 18 years as a Field Training Officer and 18 years as a remote adjunct Criminal Justice faculty member. I also supported daily users across a ten-agency electronic-crimes task force and personally processed 5,304 GB of digital evidence in 2020. In each setting, the work required me to understand what the user was trying to accomplish, explain the technology in plain language, adjust when the first approach did not work, and create documentation that made the next problem easier to solve.",
    "I have not owned SaaS renewals, ARR, NRR, Quarterly Business Reviews, or a formal enterprise customer book, and I have not operated a commercial UAS program. Those are real gaps. My relevant strengths are public-safety credibility, technology implementation from the agency side, project ownership, user instruction, workflow development, technical problem-solving, and executive-ready communication. I am developing DFR domain knowledge, can obtain FAA Part 107, and the approximately 40% travel requirement is workable for me. I am currently in Minnesota and plan to relocate to southwest Washington, both within the Northwest account footprint listed for this role.",
    "Skydio is building technology for the same public-safety professionals I spent my career working beside. I would like to help those agencies move from purchase to confident daily use.",
]


def set_run(run, size=10.25, bold=False, italic=False, color=BLACK):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.find(qn("w:rFonts"))
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.append(fonts)
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        fonts.set(qn(f"w:{attr}"), FONT)


def set_para(p, before=0, after=0, line=1.05, keep_next=False, keep_together=False):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    pf.keep_with_next = keep_next
    pf.keep_together = keep_together
    pf.widow_control = True


def prepare(kind):
    doc = Document()
    s = doc.sections[0]
    s.page_width = Inches(8.5)
    s.page_height = Inches(11)
    if kind == "resume":
        top, bottom, left, right, size = 1.46, 0.55, 0.65, 0.65, 10.25
    else:
        top, bottom, left, right, size = 1.52, 0.68, 0.78, 0.78, 10.5
    doc.styles["Normal"].font.name = FONT
    doc.styles["Normal"].font.size = Pt(size)
    doc.styles["List Bullet"].font.name = FONT
    build_navy_header(doc, body_top_margin_inches=top, body_bottom_margin_inches=bottom, body_left_margin_inches=left, body_right_margin_inches=right)
    return doc


def heading(doc, text, page_break=False):
    p = doc.add_paragraph()
    set_para(p, before=12, after=5, line=1.0, keep_next=True)
    p.paragraph_format.page_break_before = page_break
    set_run(p.add_run(text.upper()), 11, bold=True, color=STEEL)
    add_paragraph_bottom_border(p, color_hex="C9A84C", size=6)


def body(doc, text, size=10.25, after=4):
    p = doc.add_paragraph()
    set_para(p, after=after, keep_together=True)
    set_run(p.add_run(text), size)


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    set_para(p, after=2, line=1.03, keep_together=True)
    p.paragraph_format.left_indent = Inches(0.22)
    p.paragraph_format.first_line_indent = Inches(-0.14)
    for run in p.runs:
        run.text = ""
    set_run(p.add_run(text), 10.15)


def add_job(doc, item):
    title, dates, employer, bullets = item
    p = doc.add_paragraph()
    set_para(p, before=7, after=1, line=1.0, keep_next=True)
    set_run(p.add_run(title), 10.45, bold=True)
    set_run(p.add_run(" | " + dates), 9.95, bold=True, color=GRAY)
    p = doc.add_paragraph()
    set_para(p, after=3, line=1.0, keep_next=True)
    set_run(p.add_run(employer), 9.7, italic=True, color=GRAY)
    for text in bullets:
        bullet(doc, text)


def degree(doc, title, school, gpa, year):
    p = doc.add_paragraph()
    set_para(p, before=4, after=2, line=1.0, keep_together=True)
    for idx, line in enumerate((title, school, gpa, year)):
        r = p.add_run(line)
        set_run(r, 9.55, bold=(idx == 0))
        if idx < 3:
            r.add_break()


def build_resume():
    doc = prepare("resume")
    heading(doc, "Professional Summary")
    body(doc, RESUME_SUMMARY)
    heading(doc, "Customer and Implementation Capabilities")
    body(doc, CAPABILITIES)
    heading(doc, "Selected Project and Workflow Development")
    for label, text in PROJECTS:
        bullet(doc, f"{label}: {text}")
    heading(doc, "Professional Experience")
    for item in JOBS_PAGE1:
        add_job(doc, item)
    heading(doc, "Professional Experience Continued", page_break=True)
    for item in JOBS_PAGE2:
        add_job(doc, item)
    heading(doc, "Education")
    degree(doc, "Master of Arts, Police Leadership, Administration and Education", "University of St. Thomas, St. Paul, MN", "GPA: 3.94", "2005")
    degree(doc, "Bachelor of Arts, Criminal Justice, Magna Cum Laude", "St. Cloud State University, St. Cloud, MN", "GPA: 3.51", "1998")
    degree(doc, "Associate of Arts, Criminal Justice, Magna Cum Laude", "St. Cloud State University, St. Cloud, MN", "GPA: 3.50", "1996")
    heading(doc, "Selected Training and Credentials")
    body(doc, "NW3C Certified Cyber Crime Examiner, 2023 | BCA Supervision and Management series, 98 hours | Cellebrite mobile-forensics training and recertification | X-Ways Forensics training, 32 hours | FBI cell-site analysis training | Reid Technique of Interviewing and Interrogation", 9.75)
    path = OUTPUT_DIR / "Hokanson_Resume_Skydio_Customer_Success_Manager_DFR_Northwest.docx"
    doc.save(path)
    return path


def build_cover():
    doc = prepare("cover")
    p = doc.add_paragraph(); set_para(p, after=6, line=1.0); set_run(p.add_run("August 12, 2026"), 10.5)
    p = doc.add_paragraph(); set_para(p, after=8, line=1.0)
    set_run(p.add_run("Skydio"), 10.5); p.add_run().add_break(); set_run(p.add_run("Re: Customer Success Manager DFR Majors - Northwest"), 10.5, bold=True)
    for text in COVER_PARAGRAPHS:
        p = doc.add_paragraph(); set_para(p, after=8, line=1.07, keep_together=True); set_run(p.add_run(text), 10.5)
    p = doc.add_paragraph(); set_para(p, before=4, line=1.0); set_run(p.add_run("Respectfully,"), 10.5)
    p = doc.add_paragraph(); set_para(p, line=1.0); set_run(p.add_run("Troy Hokanson"), 10.5, bold=True)
    path = OUTPUT_DIR / "Hokanson_Cover_Skydio_Customer_Success_Manager_DFR_Northwest.docx"
    doc.save(path)
    return path


def render(docx_path):
    cmd = ["soffice", "--headless", "--convert-to", "pdf", "--outdir", str(OUTPUT_DIR), str(docx_path)]
    completed = subprocess.run(cmd, capture_output=True, text=True, timeout=90)
    if completed.returncode != 0:
        raise RuntimeError(completed.stderr or completed.stdout)
    return OUTPUT_DIR / (docx_path.stem + ".pdf")


def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    BUILD_LOG_DIR.mkdir(parents=True, exist_ok=True)
    resume_docx = build_resume()
    cover_docx = build_cover()
    resume_pdf = render(resume_docx)
    cover_pdf = render(cover_docx)
    scan_pdf(str(resume_pdf), doc_type="resume", profile="customer-success")
    scan_pdf(str(cover_pdf), doc_type="cover", profile="customer-success")
    print(resume_pdf)
    print(cover_pdf)


if __name__ == "__main__":
    main()
