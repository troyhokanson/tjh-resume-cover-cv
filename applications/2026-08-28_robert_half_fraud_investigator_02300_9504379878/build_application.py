#!/usr/bin/env python3

from __future__ import annotations

import json
import os
import re
import subprocess
import sys
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


REPO_ROOT = Path(__file__).resolve().parents[2]
sys.path.insert(0, str(REPO_ROOT))

required_environment = ("TROY_PHONE", "TROY_EMAIL", "TROY_LINKEDIN", "TROY_PORTFOLIO")
missing_environment = [name for name in required_environment if not os.environ.get(name, "").strip()]
if missing_environment:
    raise RuntimeError("Missing required contact environment variables: " + ", ".join(missing_environment))

import config
import docx_header
from ats_injector import ATSInjector
from docx_header import add_paragraph_bottom_border, build_navy_header


APP_DIR = Path(__file__).resolve().parent
OUTPUT_DIR = APP_DIR / "output"
BUILD_LOG_DIR = APP_DIR / "build_logs"

FONT = "EB Garamond"
BLACK = RGBColor(0x14, 0x14, 0x14)
GRAY = RGBColor(0x55, 0x55, 0x55)
STEEL = RGBColor(0x2D, 0x6A, 0x9F)

RESUME_NAME = "Hokanson_Resume_Robert_Half_Fraud_Investigator_02300-9504379878.docx"
COVER_NAME = "Hokanson_Cover_Robert_Half_Fraud_Investigator_02300-9504379878.docx"

SUMMARY = (
    "Fraud investigator with 25 years of government investigative experience, including financial fraud, asset "
    "research, financial-record analysis, evidence collection, and legal case preparation. Investigated multi-victim "
    "Business Email Compromise and occupational fraud matters from allegation through prosecution using LexisNexis "
    "Accurint, public and business records, interviews, financial documents, digital evidence, and Microsoft Excel. "
    "Holds a bachelor's degree in Criminal Justice, writes detailed reports, and presents findings to prosecutors, "
    "supervisors, command staff, and partner agencies."
)

CAPABILITIES = (
    "Confidential Fraud Investigations | Asset Searches | Financial Records and Transaction Analysis | Evidence "
    "Collection | Investigative Analytics and Pattern Analysis | LexisNexis Accurint | Public and Business Records | "
    "Investigative Research Skills and Source Corroboration | Investigative Report Writing | Legal Proceedings | "
    "Management and Prosecutor Briefings | Microsoft Excel and Word"
)

PAGE_ONE_JOBS = [
    {
        "title": "Real Estate Consultant",
        "dates": "June 2024 - March 2026",
        "employer": "Residential Real Estate | South Metro Minnesota",
        "bullets": [
            "Protected confidential client and financial information while researching property records and coordinating time-sensitive transactions. Completed $3.2 million in residential sales."
        ],
    },
    {
        "title": "Police Officer",
        "dates": "January 2022 - May 2024",
        "employer": "Lakeville Police Department | Lakeville, Minnesota",
        "bullets": [
            "Investigated allegations through interviews, records research, evidence collection, and detailed reports while making timely decisions under pressure.",
            "Reviewed reports and evidence with other officers and helped improve factual support, legal sufficiency, and written communication.",
        ],
    },
    {
        "title": "Detective / Digital Forensic Examiner",
        "dates": "June 2017 - December 2021",
        "employer": "Dakota County Electronic Crimes Task Force, assigned from Lakeville Police Department | Minnesota",
        "bullets": [
            "Served as Lakeville's investigative and digital-forensics resource in a ten-agency task force, managing confidential investigations with investigators, technical personnel, attorneys, and outside agencies.",
            "Used LexisNexis Accurint for subject location, asset identification, associate mapping, and background research, then corroborated findings through public and business records, interviews, financial documents, provider records, and digital evidence.",
            "Led a multi-victim Business Email Compromise investigation using business records, shell-company research, financial tracing, and cross-case analysis. Verified losses exceeded $360,000. The case closed with a felony conviction and written prosecutor recognition.",
            "Examined 5,304 GB of digital evidence in 2020 and converted account records, communications, and timelines into clear reports and presentations for legal and command decisions.",
        ],
    },
    {
        "title": "Detective / Electronic Crimes Unit",
        "dates": "September 2016 - June 2017",
        "employer": "Lakeville Police Department | Lakeville, Minnesota",
        "bullets": [
            "Built an investigator resource with preservation-request, subpoena, and search-warrant templates plus service-provider guidance, supporting consistent legal process, evidence collection, and case reporting."
        ],
    },
]

PAGE_TWO_JOBS = [
    {
        "title": "Police Officer / Investigator",
        "dates": "March 2010 - May 2011",
        "employer": "Lakeville Police Department | Lakeville, Minnesota",
        "bullets": [
            "Investigated occupational fraud involving approximately $80,000 in unauthorized company-card charges. Analyzed receipts, transactions, and financial records in Microsoft Excel and prepared a prosecutor-ready case package that supported a felony conviction and court-ordered restitution."
        ],
    },
    {
        "title": "Police Officer / Field Training Officer",
        "dates": "June 2011 - August 2016",
        "employer": "Lakeville Police Department | Lakeville, Minnesota",
        "bullets": [
            "Conducted investigations, documented findings, and coordinated with county community-corrections personnel in a police-probation liaison program that combined records, field observations, and compliance information."
        ],
    },
    {
        "title": "Police Officer / Field Training Officer",
        "dates": "November 1998 - February 2010",
        "employer": "Lakeville Police Department | Lakeville, Minnesota",
        "bullets": [
            "Trained and evaluated officers during 19 years of the career in interviewing, report writing, policy, evidence handling, technology, and sound decisions under pressure."
        ],
    },
    {
        "title": "Adjunct Faculty / Criminal Justice",
        "dates": "March 2007 - October 2025",
        "employer": "University of Phoenix | Remote, concurrent with sworn service",
        "bullets": [
            "Taught undergraduate Criminal Justice courses remotely for 18 years, evaluated source-supported written work, and explained investigative, legal, ethics, and evidence concepts to learners with varied experience."
        ],
    },
    {
        "title": "U.S. Army",
        "dates": "8 years 3 months",
        "employer": "Reserve, Active Duty, and Minnesota Army National Guard | Honorably Discharged",
        "bullets": [],
    },
]

EDUCATION = [
    [
        "Master of Arts, Police Leadership, Administration and Education",
        "University of St. Thomas, St. Paul, MN",
        "GPA: 3.94",
        "2005",
    ],
    [
        "Bachelor of Arts, Criminal Justice, Magna Cum Laude",
        "St. Cloud State University, St. Cloud, MN",
        "GPA: 3.51",
        "1998",
    ],
    [
        "Associate of Arts, Criminal Justice, Magna Cum Laude",
        "St. Cloud State University, St. Cloud, MN",
        "GPA: 3.50",
        "1996",
    ],
]

CREDENTIALS = (
    "Certified Fraud Examiner (CFE), actively pursuing through ACFE, 2026 | Association of Certified Fraud Examiners "
    "Member, 2026 - Present | Certified Cyber Crime Investigator, No. 4793, January 2023 | Background Investigation, "
    "St. Paul Police Department, May 2006 | Microsoft Office 365, Excel, Word, Outlook and Teams"
)

COVER_PARAGRAPHS = [
    (
        "Robert Half's Fraud Investigator assignment is a close match for the work I want to continue doing: examining "
        "suspected fraud, researching transactions and assets, weighing documentary evidence, and writing a final report "
        "that helps management act on established facts. I spent 25 years with a Minnesota government agency and handled "
        "investigations from the first allegation through interviews, records analysis, evidence collection, legal process, "
        "reporting, and prosecution."
    ),
    (
        "Two financial-fraud matters show the fit. In an occupational fraud investigation, I analyzed approximately "
        "$80,000 in unauthorized company-card charges, organized receipts and transactions in Microsoft Excel, and prepared "
        "a prosecutor-ready case package. The matter supported a felony conviction and court-ordered restitution. In a "
        "multi-victim Business Email Compromise investigation, I researched shell companies, traced financial activity, and "
        "connected records across victims. Verified losses exceeded $360,000, and the matter closed with a felony conviction "
        "and written prosecutor recognition."
    ),
    (
        "Asset searches and source verification were also regular parts of my investigative work. I used LexisNexis "
        "Accurint for subject location, asset identification, associate mapping, and background research, then corroborated "
        "the results through public and business records, interviews, financial documents, provider records, and digital "
        "evidence. I wrote detailed investigative reports, search-warrant and subpoena materials, evidence summaries, and "
        "case presentations for prosecutors, supervisors, command staff, and partner agencies. That work required careful "
        "analytics, clear verbal communication, precise written communication, and close attention to detail under pressure."
    ),
    (
        "My experience follows the posting's government-agency pathway. The underlying responsibilities are direct matches: "
        "confidential investigations, transaction and financial-record analysis, asset research, legal proceedings, complete "
        "case documentation, and findings that can withstand management and prosecutorial review. I would need to learn the "
        "client's internal systems, fraud typologies, and reporting conventions, and I am comfortable doing that within an "
        "established process."
    ),
    (
        "After retiring from sworn service, I confirmed that investigative research and analytical writing are where I do "
        "my best work. I would welcome a conversation about the client, assignment timing, expected schedule, and how my "
        "government fraud experience could support this investigation team."
    ),
]


def set_run_font(run, size: float, *, bold: bool = False, italic: bool = False, color=BLACK) -> None:
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        r_fonts.set(qn(f"w:{attr}"), FONT)


def set_paragraph(
    paragraph,
    *,
    before: float = 0,
    after: float = 0,
    line: float = 1.05,
    keep_next: bool = False,
    keep_together: bool = False,
) -> None:
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    pf.widow_control = True
    pf.keep_with_next = keep_next
    pf.keep_together = keep_together


def prepare_document(doc_type: str) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    if doc_type == "resume":
        left = right = 0.65
        bottom = 0.55
        top = 1.46
        body_size = 10.25
    else:
        left = right = 0.78
        bottom = 0.68
        top = 1.52
        body_size = 10.5

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(body_size)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.05
    list_style = doc.styles["List Bullet"]
    list_style.font.name = FONT
    list_style.font.size = Pt(body_size)

    docx_header.NAME_FONT = FONT
    docx_header.CONTACT_FONT = FONT
    docx_header.BODY_FONT = FONT
    phone_digits = re.sub(r"\D", "", config.TROY_PHONE)
    docx_header.CONTACT_PARTS = [
        (config.TROY_PHONE, f"tel:{phone_digits}"),
        (config.TROY_EMAIL, f"mailto:{config.TROY_EMAIL}"),
        (config.TROY_LINKEDIN, f"https://www.{config.TROY_LINKEDIN}"),
        ("TroyHokanson.com", config.TROY_PORTFOLIO),
    ]
    build_navy_header(
        doc,
        body_top_margin_inches=top,
        body_bottom_margin_inches=bottom,
        body_left_margin_inches=left,
        body_right_margin_inches=right,
    )
    section.different_first_page_header_footer = False
    return doc


def add_section_heading(
    doc: Document,
    text: str,
    *,
    first: bool = False,
    page_break_before: bool = False,
) -> None:
    p = doc.add_paragraph()
    set_paragraph(p, before=0 if first or page_break_before else 8, after=4, line=1.0, keep_next=True)
    p.paragraph_format.page_break_before = page_break_before
    set_run_font(p.add_run(text.upper()), 11, bold=True, color=STEEL)
    add_paragraph_bottom_border(p, color_hex="C9A84C", size=6)


def add_body_paragraph(doc: Document, text: str, *, size: float = 10.25, after: float = 4) -> None:
    p = doc.add_paragraph()
    set_paragraph(p, after=after, line=1.05, keep_together=True)
    set_run_font(p.add_run(text), size)


def add_job(doc: Document, job: dict, *, body_size: float = 10.25) -> None:
    p = doc.add_paragraph()
    set_paragraph(p, before=6, after=1, line=1.0, keep_next=True, keep_together=True)
    set_run_font(p.add_run(job["title"]), 10.5, bold=True)
    set_run_font(p.add_run(" | " + job["dates"]), 10.0, bold=True, color=GRAY)

    p = doc.add_paragraph()
    set_paragraph(p, after=2, line=1.0, keep_next=True, keep_together=True)
    set_run_font(p.add_run(job["employer"]), 9.75, italic=True, color=GRAY)

    for bullet in job["bullets"]:
        p = doc.add_paragraph(style="List Bullet")
        set_paragraph(p, after=1, line=1.03, keep_together=True)
        p.paragraph_format.left_indent = Inches(0.22)
        p.paragraph_format.first_line_indent = Inches(-0.14)
        for run in p.runs:
            run.text = ""
        set_run_font(p.add_run(bullet), body_size)


def add_degree(doc: Document, lines: list[str]) -> None:
    p = doc.add_paragraph()
    set_paragraph(p, before=4, after=1, line=1.0, keep_together=True)
    for index, line in enumerate(lines):
        set_run_font(p.add_run(line), 9.75, bold=index == 0)
        if index < len(lines) - 1:
            p.add_run().add_break()


def start_resume_continuation(doc: Document) -> None:
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    section.header_distance = Inches(0)
    section.footer_distance = Inches(0)
    section.different_first_page_header_footer = False

    header = section.header
    header.is_linked_to_previous = False
    for child in list(header._element):
        header._element.remove(child)
    header._element.append(OxmlElement("w:p"))

    footer = section.footer
    footer.is_linked_to_previous = False
    for child in list(footer._element):
        footer._element.remove(child)
    footer._element.append(OxmlElement("w:p"))


def build_resume() -> Path:
    doc = prepare_document("resume")
    doc.core_properties.title = "Troy Hokanson Resume - Robert Half Fraud Investigator"
    doc.core_properties.subject = "Application for Fraud Investigator, job 02300-9504379878"
    doc.core_properties.author = "Troy Hokanson"

    add_section_heading(doc, "Professional Summary", first=True)
    add_body_paragraph(doc, SUMMARY)
    add_section_heading(doc, "Fraud Investigation Capabilities")
    add_body_paragraph(doc, CAPABILITIES)
    add_section_heading(doc, "Professional Experience")
    for job in PAGE_ONE_JOBS:
        add_job(doc, job)

    start_resume_continuation(doc)
    add_section_heading(doc, "Additional Professional Experience", first=True)
    for job in PAGE_TWO_JOBS:
        add_job(doc, job)

    add_section_heading(doc, "Education")
    for degree in EDUCATION:
        add_degree(doc, degree)

    add_section_heading(doc, "Credentials, Training and Membership")
    add_body_paragraph(doc, CREDENTIALS, size=9.75, after=0)

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    output = OUTPUT_DIR / RESUME_NAME
    doc.save(output)
    sanitize_docx_fonts(output)
    return output


def add_cover_line(doc: Document, text: str, *, after: float = 0, bold: bool = False) -> None:
    p = doc.add_paragraph()
    set_paragraph(p, after=after, line=1.0, keep_together=True)
    set_run_font(p.add_run(text), 10.5, bold=bold)


def build_cover() -> Path:
    doc = prepare_document("cover")
    doc.core_properties.title = "Troy Hokanson Cover Letter - Robert Half Fraud Investigator"
    doc.core_properties.subject = "Application for Fraud Investigator, job 02300-9504379878"
    doc.core_properties.author = "Troy Hokanson"

    add_cover_line(doc, "August 28, 2026", after=8)
    add_cover_line(doc, "Hiring Manager")
    add_cover_line(doc, "Robert Half")
    add_cover_line(doc, "Minnetonka, Minnesota", after=8)
    add_cover_line(doc, "Dear Hiring Manager,", after=8)

    for text in COVER_PARAGRAPHS:
        p = doc.add_paragraph()
        set_paragraph(p, after=8, line=1.05, keep_together=True)
        set_run_font(p.add_run(text), 10.5)

    p = doc.add_paragraph()
    set_paragraph(p, before=2, after=0, line=1.0, keep_next=True)
    set_run_font(p.add_run("Respectfully,"), 10.5)
    p = doc.add_paragraph()
    set_paragraph(p, before=8, after=0, line=1.0, keep_together=True)
    set_run_font(p.add_run("Troy Hokanson"), 10.5)

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    output = OUTPUT_DIR / COVER_NAME
    doc.save(output)
    sanitize_docx_fonts(output)
    return output


def sanitize_docx_fonts(path: Path) -> None:
    replacements = {
        "Calibri Light": FONT,
        "Calibri": FONT,
        "Aptos Display": FONT,
        "Aptos": FONT,
        "Arial": FONT,
    }
    temporary = path.with_suffix(".fonts.docx")
    with zipfile.ZipFile(path, "r") as source, zipfile.ZipFile(temporary, "w", zipfile.ZIP_DEFLATED) as target:
        for item in source.infolist():
            data = source.read(item.filename)
            if item.filename.startswith("word/") and item.filename.endswith((".xml", ".rels")):
                text = data.decode("utf-8", errors="replace")
                for old, new in replacements.items():
                    text = re.sub(re.escape(old), new, text, flags=re.IGNORECASE)
                data = text.encode("utf-8")
            target.writestr(item, data)
    temporary.replace(path)


def inspect_docx(path: Path, doc_type: str) -> dict:
    doc = Document(path)
    section = doc.sections[0]
    expected = {
        "resume": {"left": 0.65, "right": 0.65, "bottom": 0.55, "top": 1.46},
        "cover": {"left": 0.78, "right": 0.78, "bottom": 0.68, "top": 1.52},
    }[doc_type]
    bullets = [p for p in doc.paragraphs if p.style and p.style.name == "List Bullet"]
    keep_next = sum(1 for p in doc.paragraphs if p.paragraph_format.keep_with_next)
    body_text = "\n".join(p.text for p in doc.paragraphs)
    with zipfile.ZipFile(path) as archive:
        word_xml = "\n".join(
            archive.read(name).decode("utf-8", errors="replace")
            for name in archive.namelist()
            if name.startswith("word/") and name.endswith(".xml")
        )
        rels = "\n".join(
            archive.read(name).decode("utf-8", errors="replace")
            for name in archive.namelist()
            if name.startswith("word/") and name.endswith(".rels")
        )
        header_xml = "\n".join(
            archive.read(name).decode("utf-8", errors="replace")
            for name in archive.namelist()
            if name.startswith("word/header") and name.endswith(".xml")
        )
    xml_lower = word_xml.lower()
    actual = {
        "left": round(section.left_margin.inches, 3),
        "right": round(section.right_margin.inches, 3),
        "bottom": round(section.bottom_margin.inches, 3),
        "top": round(section.top_margin.inches, 3),
    }
    checks = {
        "us_letter": round(section.page_width.inches, 2) == 8.5 and round(section.page_height.inches, 2) == 11,
        "margins_match": actual == expected,
        "approved_header_part_present": bool(header_xml.strip()) and "TroyLockedHeaderNavyBackground" in header_xml,
        "header_four_items_present": all(
            value in header_xml
            for value in (config.TROY_PHONE, config.TROY_EMAIL, config.TROY_LINKEDIN, "TroyHokanson.com")
        ),
        "header_location_absent": config.TROY_LOCATION not in header_xml,
        "continuation_header_empty": (
            len(doc.sections) == 2
            and all(not paragraph.text.strip() for paragraph in doc.sections[1].header.paragraphs)
            and round(doc.sections[1].top_margin.inches, 2) == 0.55
        ) if doc_type == "resume" else True,
        "no_body_tables": len(doc.tables) == 0,
        "real_list_style_present": len(bullets) > 0 if doc_type == "resume" else True,
        "keep_with_next_present": keep_next > 0,
        "garamond_present": "eb garamond" in xml_lower,
        "calibri_absent": "calibri" not in xml_lower,
        "aptos_absent": "aptos" not in xml_lower,
        "arial_absent": "arial" not in xml_lower,
        "contact_hyperlinks_present": "hyperlink" in rels.lower(),
        "education_exact": all(lines[0] in body_text for lines in EDUCATION) if doc_type == "resume" else True,
    }
    return {
        "path": str(path),
        "document_type": doc_type,
        "section_count": len(doc.sections),
        "paragraph_count": len(doc.paragraphs),
        "body_table_count": len(doc.tables),
        "bullet_paragraph_count": len(bullets),
        "keep_with_next_count": keep_next,
        "margins_inches": actual,
        "checks": checks,
        "passed": all(checks.values()),
    }


def save_ats_audit(resume: Path, cover: Path) -> dict:
    terms = [line.strip() for line in (APP_DIR / "ats_target_terms.txt").read_text(encoding="utf-8").splitlines() if line.strip()]
    jd_text = (APP_DIR / "job_description.md").read_text(encoding="utf-8")
    injector = ATSInjector(
        jd_text=jd_text,
        profile="analyst-intelligence",
        custom_keywords=terms,
        coverage_floor=0.85,
    )
    combined = ATSInjector.load_docx_text(str(resume)) + " " + ATSInjector.load_docx_text(str(cover))
    audit = injector.audit(combined)
    lines = [
        "ATS Audit Report",
        "Profile: analyst-intelligence",
        f"Coverage: {audit['coverage_pct']}% ({len(audit['found'])}/{audit['keywords_total']})",
        "Floor target: 85%",
        f"Meets floor: {'YES' if audit['meets_floor'] else 'NO - ACTION REQUIRED'}",
        "",
        "FOUND TERMS:",
        *[f"  + {term}" for term in sorted(audit["found"])],
        "",
        "MISSING TERMS:",
        *[f"  - {term}" for term in sorted(audit["missing"])],
    ]
    (BUILD_LOG_DIR / "ats_audit.txt").write_text("\n".join(lines) + "\n", encoding="utf-8")
    (BUILD_LOG_DIR / "ats_audit.json").write_text(json.dumps(audit, indent=2), encoding="utf-8")
    return audit


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    BUILD_LOG_DIR.mkdir(parents=True, exist_ok=True)
    if len(sys.argv) == 3 and sys.argv[1] == "--build-doc":
        if sys.argv[2] == "resume":
            build_resume()
            return
        if sys.argv[2] == "cover":
            build_cover()
            return
        raise SystemExit("--build-doc must be resume or cover")

    for document_type in ("resume", "cover"):
        subprocess.run(
            [sys.executable, str(Path(__file__).resolve()), "--build-doc", document_type],
            check=True,
            env=os.environ.copy(),
        )

    resume = OUTPUT_DIR / RESUME_NAME
    cover = OUTPUT_DIR / COVER_NAME
    structure = {
        "layout_contract": str(APP_DIR / "layout_contract.json"),
        "documents": [inspect_docx(resume, "resume"), inspect_docx(cover, "cover")],
    }
    structure["passed"] = all(item["passed"] for item in structure["documents"])
    (BUILD_LOG_DIR / "docx_structure_audit.json").write_text(json.dumps(structure, indent=2), encoding="utf-8")
    ats = save_ats_audit(resume, cover)
    print(json.dumps({"structure": structure, "ats": ats}, indent=2))
    if not structure["passed"] or not ats["meets_floor"]:
        raise SystemExit(1)


if __name__ == "__main__":
    main()
