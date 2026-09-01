#!/usr/bin/env python3

from __future__ import annotations

import json
import re
import sys
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

REPO_ROOT = Path(__file__).resolve().parents[2]
sys.path.insert(0, str(REPO_ROOT))

import docx_header
from docx_header import add_paragraph_bottom_border, build_navy_header


APP_DIR = Path(__file__).resolve().parent
OUTPUT_DIR = APP_DIR / "output"
BUILD_LOG_DIR = APP_DIR / "build_logs"
FONT = "EB Garamond"
BLACK = RGBColor(0x14, 0x14, 0x14)
GRAY = RGBColor(0x55, 0x55, 0x55)
STEEL = RGBColor(0x2D, 0x6A, 0x9F)


SUMMARY = (
    "Corporate investigations, digital-forensics, and training professional with experience conducting sensitive fraud, "
    "misconduct, threat, background, and cyber-adjacent inquiries. Medically retired after 25 years of Minnesota "
    "sworn service, including 5.5 years in electronic crimes and digital forensics. Known for careful fact-finding, "
    "defensible evidence handling, clear reports, and practical instruction for investigators, supervisors, students, "
    "and partner agencies. Minnesota private-investigation licensure preparation is in progress."
)

CAPABILITIES = (
    "General and corporate investigations | Fraud and financial crime | Digital forensics | Cyber and e-commerce inquiries | "
    "Interviews and written statements | OSINT and background research | Evidence preservation and chain of custody | "
    "Threat and workplace-safety inquiries | Affidavits and investigative reports | Litigation support | "
    "Cross-agency coordination | Training and investigative consulting"
)

PAGE1_JOBS = [
    {
        "title": "Real Estate Consultant",
        "dates": "June 2024 - March 2026",
        "employer": "Residential Real Estate | South Metro Minnesota",
        "bullets": [
            "Managed confidential information, contracts, negotiations, and time-sensitive decisions while coordinating clients and transaction partners through $3.2M in completed residential sales.",
        ],
    },
    {
        "title": "Police Officer",
        "dates": "January 2022 - May 2024",
        "employer": "Lakeville Police Department | Lakeville, Minnesota",
        "bullets": [
            "Conducted frontline inquiries involving fraud, theft, threats, harassment, interviews, digital records, and physical evidence; assessed risk, documented findings, and coordinated follow-up.",
        ],
    },
    {
        "title": "Detective / Digital Forensic Examiner",
        "dates": "June 2017 - December 2021",
        "employer": "Dakota County Electronic Crimes Task Force, assigned from Lakeville Police Department | Minnesota",
        "bullets": [
            "Served as Lakeville Police Department's representative and digital-forensics subject-matter resource in a ten-agency task force, conducting examinations for Lakeville and partner-agency investigations.",
            "Processed 5,304 GB of digital evidence in 2020 using Cellebrite, GrayKey, X-Ways Forensics, Magnet AXIOM, and AccessData Forensic Toolkit; converted technical findings into clear reports and briefings.",
            "Led a multi-victim Business Email Compromise investigation involving shell companies and interstate transfers that documented verified losses exceeding $360,000, resulted in a felony conviction, and earned written recognition from an Assistant Dakota County Attorney.",
        ],
    },
    {
        "title": "Detective / Electronic Crimes Unit",
        "dates": "September 2016 - June 2017",
        "employer": "Lakeville Police Department | Lakeville, Minnesota",
        "bullets": [
            "Acquired and configured the unit's initial Cellebrite UFED and built an investigator resource package with preservation, subpoena, search-warrant, and service-provider guidance.",
        ],
    },
]

PAGE2_JOBS = [
    {
        "title": "Police Officer / Investigator / Field Training Officer",
        "dates": "November 1998 - August 2016",
        "employer": "Lakeville Police Department | Lakeville, Minnesota",
        "bullets": [
            "Led an occupational-fraud investigation involving approximately $80,000 in unauthorized company-card charges, compiled receipts and transaction evidence into an Excel summary, and delivered a complete case package that resulted in a felony conviction and court-ordered restitution.",
            "Conducted general, property, fraud, threat, and safety investigations through interviews, records review, surveillance evidence, warrants, interagency coordination, and written findings.",
            "Served 19 years as a Field Training Officer, teaching policy, interviewing, investigations, technology, documentation, communication, and sound decision-making.",
        ],
    },
    {
        "title": "Adjunct Faculty / Criminal Justice",
        "dates": "March 2007 - October 2025",
        "employer": "University of Phoenix | Remote, concurrent with sworn service",
        "bullets": [
            "Taught undergraduate Criminal Justice courses remotely for 18 years, explaining investigative, legal, ethical, and technical material to adult learners.",
            "Received the Phoenix500 Faculty Excellence Award in 2020 and 2021 and a Faculty of the Year nomination in 2021.",
        ],
    },
    {
        "title": "U.S. Army",
        "dates": "8 years 3 months",
        "employer": "Reserve, Active Duty, and Minnesota Army National Guard | Honorably Discharged",
        "bullets": [],
    },
]

TRAINING = [
    "Minnesota Board-approved 12-hour private-detective pre-assignment training and six hours of continuing education, completed 2026; American Heart Association CPR/AED training, completed 2026.",
    "Certified Fraud Examiner (CFE), actively pursuing through ACFE, 2026; Certified Cyber Crime Investigator (CCCI) No. 4793, January 2023.",
    "Cellebrite CCLO and CCPA, 2016, recertified 2020; Reid Technique of Interviewing and Interrogation.",
    "FBI Cell-Site Analysis and Location Data; NW3C Cybercrime Investigation.",
]


def set_run_font(run, size: float, *, bold: bool = False, italic: bool = False, color=BLACK) -> None:
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        r_fonts.set(qn(f"w:{attr}"), FONT)


def set_paragraph(paragraph, *, before=0, after=0, line=1.05, keep_next=False, keep_together=False) -> None:
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    pf.widow_control = True
    pf.keep_with_next = keep_next
    pf.keep_together = keep_together


def prepare_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(10.25)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.05

    list_style = doc.styles["List Bullet"]
    list_style.font.name = FONT
    list_style.font.size = Pt(10.25)

    docx_header.NAME_FONT = FONT
    docx_header.CONTACT_FONT = FONT
    docx_header.BODY_FONT = FONT
    build_navy_header(
        doc,
        body_top_margin_inches=1.46,
        body_bottom_margin_inches=0.55,
        body_left_margin_inches=0.65,
        body_right_margin_inches=0.65,
    )
    return doc


def add_section_heading(doc: Document, text: str, *, page_break_before: bool = False) -> None:
    p = doc.add_paragraph()
    set_paragraph(p, before=14, after=6, line=1.0, keep_next=True)
    p.paragraph_format.page_break_before = page_break_before
    run = p.add_run(text.upper())
    set_run_font(run, 11, bold=True, color=STEEL)
    add_paragraph_bottom_border(p, color_hex="C9A84C", size=6)


def add_body_paragraph(doc: Document, text: str, *, after=4) -> None:
    p = doc.add_paragraph()
    set_paragraph(p, before=0, after=after, line=1.05, keep_together=True)
    set_run_font(p.add_run(text), 10.25)


def add_bullet(doc: Document, text: str, *, size=10.25, keep_next=False) -> None:
    p = doc.add_paragraph(style="List Bullet")
    set_paragraph(p, before=0, after=2, line=1.05, keep_together=True, keep_next=keep_next)
    p.paragraph_format.left_indent = Inches(0.22)
    p.paragraph_format.first_line_indent = Inches(-0.14)
    for run in p.runs:
        run.text = ""
    set_run_font(p.add_run(text), size)


def add_job(doc: Document, job: dict) -> None:
    p = doc.add_paragraph()
    set_paragraph(p, before=8, after=2, line=1.0, keep_next=True, keep_together=True)
    set_run_font(p.add_run(job["title"]), 10.5, bold=True)
    set_run_font(p.add_run(" | " + job["dates"]), 10.0, bold=True, color=GRAY)

    p = doc.add_paragraph()
    set_paragraph(p, before=0, after=4, line=1.0, keep_next=bool(job["bullets"]), keep_together=True)
    set_run_font(p.add_run(job["employer"]), 9.75, italic=True, color=GRAY)

    for index, bullet in enumerate(job["bullets"]):
        add_bullet(doc, bullet, keep_next=False)


def add_degree(doc: Document, lines: list[str]) -> None:
    p = doc.add_paragraph()
    set_paragraph(p, before=5, after=2, line=1.0, keep_together=True)
    for index, line in enumerate(lines):
        run = p.add_run(line)
        set_run_font(run, 9.75, bold=index == 0)
        if index < len(lines) - 1:
            run.add_break()


def sanitize_docx_fonts(path: Path) -> None:
    replacements = {
        "Calibri Light": FONT,
        "Calibri": FONT,
        "Aptos Display": FONT,
        "Aptos": FONT,
        "Arial": FONT,
    }
    temporary = path.with_suffix(".fonts.docx")
    with zipfile.ZipFile(path, "r") as source, zipfile.ZipFile(temporary, "w", zipfile.ZIP_DEFLATED) as target:
        for item in source.infolist():
            data = source.read(item.filename)
            if item.filename.startswith("word/") and item.filename.endswith((".xml", ".rels")):
                text = data.decode("utf-8", errors="replace")
                for old, new in replacements.items():
                    text = re.sub(re.escape(old), new, text, flags=re.IGNORECASE)
                data = text.encode("utf-8")
            target.writestr(item, data)
    temporary.replace(path)


def build_resume() -> Path:
    doc = prepare_document()
    add_section_heading(doc, "Professional Summary")
    add_body_paragraph(doc, SUMMARY)
    add_section_heading(doc, "Investigative Capabilities")
    add_body_paragraph(doc, CAPABILITIES)
    add_section_heading(doc, "Professional Experience")
    for job in PAGE1_JOBS:
        add_job(doc, job)

    add_section_heading(doc, "Additional Investigative and Training Experience")
    for job in PAGE2_JOBS:
        add_job(doc, job)

    add_section_heading(doc, "Private Investigation and Professional Training")
    for item in TRAINING:
        add_bullet(doc, item, size=9.75)

    add_section_heading(doc, "Education")
    add_degree(doc, [
        "Master of Arts, Police Leadership, Administration and Education",
        "University of St. Thomas, St. Paul, MN",
        "GPA: 3.94",
        "2005",
    ])
    add_degree(doc, [
        "Bachelor of Arts, Criminal Justice, Magna Cum Laude",
        "St. Cloud State University, St. Cloud, MN",
        "GPA: 3.51",
        "1998",
    ])
    add_degree(doc, [
        "Associate of Arts, Criminal Justice, Magna Cum Laude",
        "St. Cloud State University, St. Cloud, MN",
        "GPA: 3.50",
        "1996",
    ])

    output = OUTPUT_DIR / "Hokanson_Resume_360_Security_Services_Investigator_Consultant.docx"
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    sanitize_docx_fonts(output)
    return output


def inspect_docx(path: Path) -> dict:
    doc = Document(path)
    section = doc.sections[0]
    bullet_paragraphs = [p for p in doc.paragraphs if p.style and p.style.name == "List Bullet"]
    keep_next = sum(1 for p in doc.paragraphs if p.paragraph_format.keep_with_next)
    with zipfile.ZipFile(path) as archive:
        xml = "\n".join(
            archive.read(name).decode("utf-8", errors="replace")
            for name in archive.namelist()
            if name.startswith("word/") and name.endswith(".xml")
        ).lower()
        rels = "\n".join(
            archive.read(name).decode("utf-8", errors="replace")
            for name in archive.namelist()
            if name.startswith("word/") and name.endswith(".rels")
        )

    actual = {
        "left": round(section.left_margin.inches, 3),
        "right": round(section.right_margin.inches, 3),
        "bottom": round(section.bottom_margin.inches, 3),
        "top": round(section.top_margin.inches, 3),
    }
    expected = {"left": 0.65, "right": 0.65, "bottom": 0.55, "top": 1.46}
    checks = {
        "us_letter": round(section.page_width.inches, 2) == 8.5 and round(section.page_height.inches, 2) == 11,
        "margins_match": actual == expected,
        "header_part_present": bool(doc.sections[0].header.paragraphs),
        "real_list_style_present": len(bullet_paragraphs) > 0,
        "keep_with_next_present": keep_next > 0,
        "garamond_present": "eb garamond" in xml,
        "calibri_absent": "calibri" not in xml,
        "aptos_absent": "aptos" not in xml,
        "arial_absent": "arial" not in xml,
        "contact_hyperlinks_present": "hyperlink" in rels.lower(),
    }
    return {
        "path": str(path),
        "document_type": "resume",
        "section_count": len(doc.sections),
        "paragraph_count": len(doc.paragraphs),
        "bullet_paragraph_count": len(bullet_paragraphs),
        "keep_with_next_count": keep_next,
        "margins_inches": actual,
        "checks": checks,
        "passed": all(checks.values()),
    }


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    BUILD_LOG_DIR.mkdir(parents=True, exist_ok=True)
    resume = build_resume()
    report = {
        "layout_contract": str(APP_DIR / "layout_contract.json"),
        "documents": [inspect_docx(resume)],
    }
    report["passed"] = all(item["passed"] for item in report["documents"])
    (BUILD_LOG_DIR / "docx_structure_audit.json").write_text(json.dumps(report, indent=2), encoding="utf-8")
    print(json.dumps(report, indent=2))
    if not report["passed"]:
        raise SystemExit(1)


if __name__ == "__main__":
    main()
