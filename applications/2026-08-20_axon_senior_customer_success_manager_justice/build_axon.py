#!/usr/bin/env python3

from __future__ import annotations

import json
import re
import sys
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

REPO_ROOT = Path(__file__).resolve().parents[2]
sys.path.insert(0, str(REPO_ROOT))

import docx_header
from docx_header import add_paragraph_bottom_border, build_navy_header


APP_DIR = Path(__file__).resolve().parent
OUTPUT_DIR = APP_DIR / "output"
BUILD_LOG_DIR = APP_DIR / "build_logs"
FONT = "EB Garamond"
BLACK = RGBColor(0x14, 0x14, 0x14)
GRAY = RGBColor(0x55, 0x55, 0x55)
STEEL = RGBColor(0x2D, 0x6A, 0x9F)


RESUME = {
    "summary": (
        "Justice-sector and public-safety technology professional with 25 years of sworn service, "
        "5.5 years in electronic crimes and digital forensics, and 18 years of remote college teaching. "
        "Worked across law enforcement, prosecutors, courts, technical personnel, and partner agencies to "
        "manage digital evidence, legal process, case priorities, training, and time-sensitive escalations. "
        "Direct Axon end-user experience includes Body 3, Fleet 2, and Evidence workflows."
    ),
    "capabilities": (
        "Justice stakeholder engagement | Discovery and disclosure workflows | Digital evidence management | "
        "Case and project coordination | Customer adoption and enablement | Success planning and progress tracking | "
        "Escalation management | Implementation support | Training and presentations | Executive and technical communication | "
        "Axon Body 3, Fleet 2, and Evidence | Cellebrite, Magnet AXIOM, X-Ways, FTK, and GrayKey"
    ),
    "training": [
        "NW3C Certified Cybercrime Investigator (CCCI), No. 4793, January 2023.",
        "BCA Law Enforcement Supervision & Management Program, 98 hours, including supervision, mentoring and coaching, influence, legal latitude, stress management, and ethics.",
        "University of Phoenix Certified Advanced Facilitator.",
        "Cellebrite Certified Logical Operator (CCLO) and Certified Physical Analyst (CCPA), 2016; recertified 2018, 2020.",
        "BCA Forensic Science Partners, 58 hours, 2013; BCA Certified Crime Scene Technician.",
        "Crisis Intervention Team (CIT) training, 8 hours, including classroom instruction and scenario-based de-escalation practice; Conflict Management & Mediation Training, 2 hours, 2022; Implicit Bias / Community Diversity, 4 hours, 2022.",
    ],
    "page1_jobs": [
        {
            "title": "Generalist Expert",
            "dates": "August 2026 - Present",
            "employer": "Mercor | Remote independent contractor",
            "bullets": [
                "Complete paid AI project work involving close reading, source validation, written analysis, and adherence to confidential task requirements.",
            ],
        },
        {
            "title": "Real Estate Consultant",
            "dates": "June 2024 - March 2026",
            "employer": "eXp Realty / KW Select | South Metro Minnesota",
            "bullets": [
                "Managed client relationships from initial consultation through negotiation, inspection, financing, title, and closing; completed $3.2M in residential sales during the transition from law enforcement.",
                "Coordinated clients and transaction partners through deadlines, competing priorities, documentation issues, and escalations while keeping each party informed of decisions and next steps.",
            ],
        },
        {
            "title": "Police Officer",
            "dates": "January 2022 - May 2024",
            "employer": "Lakeville Police Department | Lakeville, Minnesota",
            "bullets": [
                "Used Axon Body 3, Axon Fleet 2, and Evidence workflows in daily patrol operations and relied on the resulting video and digital records for reports, investigations, disclosure, and courtroom preparation.",
                "Helped officers and supervisors resolve policy, evidence, documentation, and technology questions in time-sensitive operating conditions.",
            ],
        },
        {
            "title": "Detective / Digital Forensic Examiner",
            "dates": "June 2017 - December 2021",
            "employer": "Dakota County Electronic Crimes Task Force, assigned from Lakeville Police Department | Minnesota",
            "bullets": [
                "Served as Lakeville's representative and digital-forensics subject-matter resource in a ten-agency task force, coordinating examination priorities, legal-process needs, and technical guidance across investigators, supervisors, prosecutors, and partner agencies.",
                "Personally processed 5,304 GB of digital evidence in 2020 using Cellebrite, GrayKey, Magnet AXIOM, X-Ways, and related platforms; maintained defensible records and translated findings for charging, discovery, and case decisions.",
            ],
        },
    ],
    "page2_jobs": [
        {
            "title": "Detective / Electronic Crimes Unit",
            "dates": "September 2016 - June 2017",
            "employer": "Lakeville Police Department | Lakeville, Minnesota",
            "bullets": [
                "Used the unit's existing Cellebrite UFED in electronic-crimes investigations and helped investigators integrate mobile-device evidence into case and disclosure workflows.",
                "Built a structured investigator resource with preservation, subpoena, search-warrant, and service-provider templates that reduced the learning curve and created a more consistent legal-process workflow.",
            ],
        },
        {
            "title": "Police Officer / Field Training Officer",
            "dates": "November 1998 - August 2016",
            "employer": "Lakeville Police Department | Lakeville, Minnesota",
            "bullets": [
                "Served 19 years as a Field Training Officer, coaching officers through policy, technology, report writing, evidence handling, courtroom preparation, and sound decisions in changing conditions.",
                "Led an agency-side ALPR program from 2007 to 2010, helped secure a $40,000 Target + Blue grant, and coordinated Genetec AutoVu, BCA CJIS, city IT, leadership, and daily users around implementation and adoption.",
                "Selected as one of two officers to launch the 2012 Probation Liaison Officer Program with Dakota County Community Corrections; strengthened communication, supported joint compliance and warrant work, and earned written recognition from department leadership and a county probation officer.",
            ],
        },
        {
            "title": "Adjunct Faculty / Criminal Justice",
            "dates": "March 2007 - October 2025",
            "employer": "University of Phoenix | Remote, concurrent with sworn service",
            "bullets": [
                "Taught undergraduate Criminal Justice courses remotely for 18 years, converting legal, investigative, ethical, and technical subjects into structured instruction for adult learners.",
                "Received Phoenix500 Faculty Excellence recognition in 2020 and 2021 and the John Sperling Distinguished Faculty Award in 2024.",
            ],
        },
        {
            "title": "U.S. Army",
            "dates": "8 years 3 months",
            "employer": "Reserve, Active Duty, and Minnesota Army National Guard | Honorably Discharged",
            "bullets": [],
        },
    ],
}


COVER_PARAGRAPHS = [
    (
        "Axon's Justice Senior Customer Success Manager role is a direct match for the work I know from the "
        "end-user side of the workflow. I used Axon Body 3, Fleet 2, and Evidence as an officer, relied on the "
        "resulting evidence as an investigator, and understand the policy, activation, disclosure, and courtroom "
        "consequences when connected evidence systems work well or create friction."
    ),
    (
        "My justice-sector experience was built across 25 years in Minnesota law enforcement, including 5.5 years "
        "in electronic crimes and digital forensics. As Lakeville's representative in a ten-agency task force, I "
        "coordinated examination priorities and legal-process needs with investigators, supervisors, prosecutors, "
        "technical personnel, and partner agencies. I personally processed 5,304 GB of digital evidence in 2020 and "
        "translated technical findings into reports and briefings that supported charging, discovery, and case decisions."
    ),
    (
        "I also know what it takes to make a technical workflow usable. I used Cellebrite UFED in electronic-crimes "
        "investigations and built an investigator resource with preservation, subpoena, search-warrant, and provider templates. "
        "Earlier, I led an agency-side ALPR program that brought Genetec, state CJIS personnel, city IT, department "
        "leadership, and daily users together around implementation and adoption. Both projects required clear ownership, "
        "steady follow-through, and practical answers when technical and operational priorities did not line up cleanly. "
        "That experience is backed by 98 hours in the BCA Law Enforcement Supervision & Management Program and the "
        "NW3C Certified Cybercrime Investigator credential."
    ),
    (
        "I have not worked inside a prosecutor or public defender office, and I would not represent my background that "
        "way. I was, however, selected as one of two officers to launch a Probation Liaison Officer Program with Dakota "
        "County Community Corrections. Written recognition from both department leadership and a county probation officer "
        "confirms that the work improved communication and operational follow-through. In the county task-force setting, "
        "routine access to County Attorney personnel made prosecutor coordination part of the daily case workflow."
    ),
    (
        "I currently live in Minnesota and plan to relocate to southwest Washington in 2027. I can work West Coast hours "
        "immediately and am comfortable with the role's 30-35 percent travel expectation. Axon Justice would let me apply "
        "field-tested evidence judgment, product familiarity, and disciplined customer follow-through to the prosecutors "
        "and public defenders responsible for moving cases fairly and on time."
    ),
]


def set_run_font(run, size: float, *, bold: bool = False, italic: bool = False, color=BLACK) -> None:
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        r_fonts.set(qn(f"w:{attr}"), FONT)


def set_paragraph(paragraph, *, before=0, after=0, line=1.05, keep_next=False, keep_together=False) -> None:
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    pf.widow_control = True
    pf.keep_with_next = keep_next
    pf.keep_together = keep_together


def prepare_document(doc_type: str) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)

    if doc_type == "resume":
        left = right = 0.65
        bottom = 0.55
        top = 1.46
        body_size = 10.25
    else:
        left = right = 0.78
        bottom = 0.68
        top = 1.52
        body_size = 10.5

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(body_size)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.05

    list_style = doc.styles["List Bullet"]
    list_style.font.name = FONT
    list_style.font.size = Pt(body_size)

    docx_header.NAME_FONT = FONT
    docx_header.CONTACT_FONT = FONT
    docx_header.BODY_FONT = FONT
    # Locked application header: phone, canonical email, LinkedIn, and portfolio.
    # Location remains excluded from the contact row.
    docx_header.CONTACT_PARTS = [
        item
        for item in docx_header.CONTACT_PARTS
        if item[0] != docx_header.TROY_LOCATION
    ]
    build_navy_header(
        doc,
        body_top_margin_inches=top,
        body_bottom_margin_inches=bottom,
        body_left_margin_inches=left,
        body_right_margin_inches=right,
    )
    return doc


def add_section_heading(doc: Document, text: str, *, page_break_before: bool = False) -> None:
    p = doc.add_paragraph()
    set_paragraph(p, before=14, after=6, line=1.0, keep_next=True)
    p.paragraph_format.page_break_before = page_break_before
    run = p.add_run(text.upper())
    set_run_font(run, 11, bold=True, color=STEEL)
    add_paragraph_bottom_border(p, color_hex="C9A84C", size=6)


def add_body_paragraph(doc: Document, text: str, *, size=10.25, after=4) -> None:
    p = doc.add_paragraph()
    set_paragraph(p, before=0, after=after, line=1.05, keep_together=True)
    set_run_font(p.add_run(text), size)


def add_job(doc: Document, job: dict, *, compact: bool = False) -> None:
    p = doc.add_paragraph()
    set_paragraph(p, before=6 if compact else 8, after=2, line=1.0, keep_next=True, keep_together=True)
    set_run_font(p.add_run(job["title"]), 10.25 if compact else 10.5, bold=True)
    set_run_font(p.add_run(" | " + job["dates"]), 9.75 if compact else 10.0, bold=True, color=GRAY)

    p = doc.add_paragraph()
    set_paragraph(p, before=0, after=3 if compact else 4, line=1.0, keep_next=True, keep_together=True)
    set_run_font(p.add_run(job["employer"]), 9.5 if compact else 9.75, italic=True, color=GRAY)

    for index, bullet in enumerate(job["bullets"]):
        p = doc.add_paragraph(style="List Bullet")
        set_paragraph(
            p,
            before=0,
            after=1 if compact else 2,
            line=1.0 if compact else 1.05,
            keep_together=True,
            keep_next=False,
        )
        p.paragraph_format.left_indent = Inches(0.22)
        p.paragraph_format.first_line_indent = Inches(-0.14)
        for run in p.runs:
            run.text = ""
        set_run_font(p.add_run(bullet), 9.9 if compact else 10.25)


def add_degree(doc: Document, lines: list[str]) -> None:
    p = doc.add_paragraph()
    set_paragraph(p, before=3, after=1, line=1.0, keep_together=True)
    set_run_font(p.add_run(lines[0]), 9.35, bold=True)
    if len(lines) > 1:
        set_run_font(p.add_run(" | " + lines[1]), 9.35)


def add_training_item(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    set_paragraph(p, before=0, after=1, line=1.0, keep_together=True)
    p.paragraph_format.left_indent = Inches(0.22)
    p.paragraph_format.first_line_indent = Inches(-0.14)
    for run in p.runs:
        run.text = ""
    set_run_font(p.add_run(text), 9.5)


def build_resume() -> Path:
    doc = prepare_document("resume")
    add_section_heading(doc, "Professional Summary")
    add_body_paragraph(doc, RESUME["summary"], after=4)
    add_section_heading(doc, "Customer and Technical Capabilities")
    add_body_paragraph(doc, RESUME["capabilities"], after=4)
    add_section_heading(doc, "Professional Experience")
    for job in RESUME["page1_jobs"]:
        add_job(doc, job)

    add_section_heading(doc, "Additional Experience", page_break_before=True)
    for job in RESUME["page2_jobs"]:
        add_job(doc, job, compact=True)

    add_section_heading(doc, "Training and Certifications")
    for item in RESUME["training"]:
        add_training_item(doc, item)

    add_section_heading(doc, "Education")
    add_degree(doc, [
        "Master of Arts, Police Leadership, Administration and Education",
        "University of St. Thomas, St. Paul, MN | GPA: 3.94 | 2005",
    ])
    add_degree(doc, [
        "Bachelor of Arts, Criminal Justice, Magna Cum Laude",
        "St. Cloud State University, St. Cloud, MN | GPA: 3.51 | 1998",
    ])
    add_degree(doc, [
        "Associate of Arts, Criminal Justice, Magna Cum Laude",
        "St. Cloud State University, St. Cloud, MN | GPA: 3.50 | 1996",
    ])

    output = OUTPUT_DIR / "Hokanson_Resume_Axon_Sr_Customer_Success_Manager_Justice.docx"
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    sanitize_docx_fonts(output)
    return output


def add_cover_line(doc: Document, text: str, *, after=0, bold=False) -> None:
    p = doc.add_paragraph()
    set_paragraph(p, before=0, after=after, line=1.0, keep_together=True)
    set_run_font(p.add_run(text), 10.5, bold=bold)


def build_cover() -> Path:
    doc = prepare_document("cover")
    add_cover_line(doc, "August 20, 2026", after=8)
    add_cover_line(doc, "Hiring Manager")
    add_cover_line(doc, "Axon Enterprise, Inc.", after=9)
    add_cover_line(doc, "Dear Hiring Manager,", after=9)

    for paragraph in COVER_PARAGRAPHS:
        p = doc.add_paragraph()
        set_paragraph(p, before=0, after=8, line=1.05, keep_together=True)
        set_run_font(p.add_run(paragraph), 10.5)

    p = doc.add_paragraph()
    set_paragraph(p, before=5, after=0, line=1.0, keep_next=True)
    set_run_font(p.add_run("Respectfully,"), 10.5)
    p = doc.add_paragraph()
    set_paragraph(p, before=38, after=0, line=1.0, keep_together=True)
    set_run_font(p.add_run("Troy Hokanson"), 10.5)

    output = OUTPUT_DIR / "Hokanson_Cover_Axon_Sr_Customer_Success_Manager_Justice.docx"
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    sanitize_docx_fonts(output)
    return output


def sanitize_docx_fonts(path: Path) -> None:
    replacements = {
        "Calibri Light": FONT,
        "Calibri": FONT,
        "Aptos Display": FONT,
        "Aptos": FONT,
        "Arial": FONT,
    }
    temporary = path.with_suffix(".fonts.docx")
    with zipfile.ZipFile(path, "r") as source, zipfile.ZipFile(temporary, "w", zipfile.ZIP_DEFLATED) as target:
        for item in source.infolist():
            data = source.read(item.filename)
            if item.filename.startswith("word/") and item.filename.endswith((".xml", ".rels")):
                text = data.decode("utf-8", errors="replace")
                for old, new in replacements.items():
                    text = re.sub(re.escape(old), new, text, flags=re.IGNORECASE)
                data = text.encode("utf-8")
            target.writestr(item, data)
    temporary.replace(path)


def inspect_docx(path: Path, doc_type: str) -> dict:
    doc = Document(path)
    section = doc.sections[0]
    expected = {
        "resume": {"left": 0.65, "right": 0.65, "bottom": 0.55, "top": 1.46},
        "cover": {"left": 0.78, "right": 0.78, "bottom": 0.68, "top": 1.52},
    }[doc_type]

    bullet_paragraphs = [p for p in doc.paragraphs if p.style and p.style.name == "List Bullet"]
    keep_next = sum(1 for p in doc.paragraphs if p.paragraph_format.keep_with_next)
    with zipfile.ZipFile(path) as archive:
        xml = "\n".join(
            archive.read(name).decode("utf-8", errors="replace")
            for name in archive.namelist()
            if name.startswith("word/") and name.endswith(".xml")
        ).lower()
        rels = "\n".join(
            archive.read(name).decode("utf-8", errors="replace")
            for name in archive.namelist()
            if name.startswith("word/") and name.endswith(".rels")
        )

    actual = {
        "left": round(section.left_margin.inches, 3),
        "right": round(section.right_margin.inches, 3),
        "bottom": round(section.bottom_margin.inches, 3),
        "top": round(section.top_margin.inches, 3),
    }
    checks = {
        "us_letter": round(section.page_width.inches, 2) == 8.5 and round(section.page_height.inches, 2) == 11,
        "margins_match": actual == expected,
        "header_part_present": bool(doc.sections[0].header.paragraphs),
        "real_list_style_present": len(bullet_paragraphs) > 0 if doc_type == "resume" else True,
        "keep_with_next_present": keep_next > 0,
        "garamond_present": "eb garamond" in xml,
        "calibri_absent": "calibri" not in xml,
        "aptos_absent": "aptos" not in xml,
        "arial_absent": "arial" not in xml,
        "contact_hyperlinks_present": "hyperlink" in rels.lower(),
    }
    return {
        "path": str(path),
        "document_type": doc_type,
        "section_count": len(doc.sections),
        "paragraph_count": len(doc.paragraphs),
        "bullet_paragraph_count": len(bullet_paragraphs),
        "keep_with_next_count": keep_next,
        "margins_inches": actual,
        "checks": checks,
        "passed": all(checks.values()),
    }


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    BUILD_LOG_DIR.mkdir(parents=True, exist_ok=True)
    resume = build_resume()
    cover = build_cover()
    report = {
        "layout_contract": str(APP_DIR / "layout_contract.json"),
        "documents": [inspect_docx(resume, "resume"), inspect_docx(cover, "cover")],
    }
    report["passed"] = all(item["passed"] for item in report["documents"])
    (BUILD_LOG_DIR / "docx_structure_audit.json").write_text(
        json.dumps(report, indent=2), encoding="utf-8"
    )
    print(json.dumps(report, indent=2))
    if not report["passed"]:
        raise SystemExit(1)


if __name__ == "__main__":
    main()
