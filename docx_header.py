"""
Troy Hokanson - Locked DOCX Header Module
========================================

Single source of truth for the navy/gold/white header used on every DOCX
resume, cover letter, CV, and application document bearing Troy's name.

Locked header requirements:
- Full-bleed navy #0D1B2A bar in the Word header part, never in the body.
- "Troy Hokanson" in white #FFFFFF Garamond-family bold, centered, 26 pt.
- Thin gold #C9A84C rule directly beneath the name.
- Gold Garamond contact row beneath the rule.
- No subtitle or role title between name and contact row.
- No footer by default.

Do not hand-roll the header in build scripts. Import build_navy_header(doc).
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement, parse_xml
from config import (
    TROY_NAME,
    TROY_PHONE,
    TROY_EMAIL,
    TROY_LOCATION,
    TROY_LINKEDIN,
    TROY_PORTFOLIO,
)

NAVY = RGBColor(0x0D, 0x1B, 0x2A)
GOLD = RGBColor(0xC9, 0xA8, 0x4C)
STEEL = RGBColor(0x2D, 0x6A, 0x9F)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x14, 0x14, 0x14)
GRAY = RGBColor(0x55, 0x55, 0x55)

BODY_FONT = "Calibri"
NAME_FONT = "Garamond"
CONTACT_FONT = "Garamond"
NAME = TROY_NAME or "Troy Hokanson"

_phone_digits = TROY_PHONE.replace(".", "").replace("-", "").replace(" ", "")
CONTACT_PARTS = [
    (TROY_LOCATION, None),
    *([(TROY_PHONE, f"tel:{_phone_digits}")] if TROY_PHONE else []),
    (TROY_EMAIL, f"mailto:{TROY_EMAIL}"),
    (
        TROY_LINKEDIN,
        f"https://www.{TROY_LINKEDIN}"
        if TROY_LINKEDIN and not TROY_LINKEDIN.startswith("http")
        else TROY_LINKEDIN,
    ),
    ("TroyHokanson.com", TROY_PORTFOLIO),
]


def _rgb_to_hex(color: RGBColor) -> str:
    return f"{color[0]:02X}{color[1]:02X}{color[2]:02X}"


def shade_cell(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    for old in tc_pr.findall(qn("w:shd")):
        tc_pr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def shade_paragraph(paragraph, color_hex):
    p_pr = paragraph._p.get_or_add_pPr()
    for old in p_pr.findall(qn("w:shd")):
        p_pr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    p_pr.append(shd)


def add_header_background_shape(paragraph, color_hex="0D1B2A"):
    run = paragraph.add_run()
    xml = f"""
    <w:pict xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
            xmlns:v="urn:schemas-microsoft-com:vml"
            xmlns:o="urn:schemas-microsoft-com:office:office">
      <v:rect id="TroyLockedHeaderNavyBackground"
              fillcolor="#{color_hex}"
              stroked="f"
              style="position:absolute;margin-left:-90pt;margin-top:-28pt;width:800pt;height:112pt;z-index:-251654144;mso-position-horizontal:absolute;mso-position-horizontal-relative:page;mso-position-vertical:absolute;mso-position-vertical-relative:page">
        <v:fill color="#{color_hex}"/>
      </v:rect>
    </w:pict>
    """
    run._r.append(parse_xml(xml))
    return run


def set_cell_margins(cell, top=100, bottom=100, left=140, right=140):
    tc_pr = cell._tc.get_or_add_tcPr()
    for old in tc_pr.findall(qn("w:tcMar")):
        tc_pr.remove(old)
    tc_mar = OxmlElement("w:tcMar")
    for margin_name, value in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        node = OxmlElement(f"w:{margin_name}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        tc_mar.append(node)
    tc_pr.append(tc_mar)


def remove_cell_borders(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    for old in tc_pr.findall(qn("w:tcBorders")):
        tc_pr.remove(old)
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "nil")
        tc_borders.append(border)
    tc_pr.append(tc_borders)


def add_paragraph_bottom_border(paragraph, color_hex="C9A84C", size=6, space=1):
    p_pr = paragraph._p.get_or_add_pPr()
    for old in p_pr.findall(qn("w:pBdr")):
        p_pr.remove(old)
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color_hex)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def set_run(run, *, font=BODY_FONT, size=10.5, bold=False, italic=False, color=BLACK):
    run.font.name = font
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        r_fonts.set(qn(f"w:{attr}"), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def add_hyperlink(paragraph, text, url, *, color=None, font=BODY_FONT, size=10, bold=False):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    r_fonts = OxmlElement("w:rFonts")
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        r_fonts.set(qn(f"w:{attr}"), font)
    r_pr.append(r_fonts)

    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size * 2)))
    r_pr.append(sz)
    sz_cs = OxmlElement("w:szCs")
    sz_cs.set(qn("w:val"), str(int(size * 2)))
    r_pr.append(sz_cs)

    if bold:
        r_pr.append(OxmlElement("w:b"))
    if color is not None:
        color_node = OxmlElement("w:color")
        color_node.set(qn("w:val"), _rgb_to_hex(color))
        r_pr.append(color_node)

    new_run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    text_node.set(qn("xml:space"), "preserve")
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def set_paragraph_format(paragraph, *, before=0, after=0, line=1.15):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def build_navy_header(
    doc,
    *,
    body_top_margin_inches=1.55,
    body_bottom_margin_inches=0.55,
    body_left_margin_inches=0.6,
    body_right_margin_inches=0.6,
):
    """Build the full-bleed navy/gold/white header in the Word header part."""
    section = doc.sections[0]
    section.top_margin = Inches(body_top_margin_inches)
    section.bottom_margin = Inches(body_bottom_margin_inches)
    section.left_margin = Inches(body_left_margin_inches)
    section.right_margin = Inches(body_right_margin_inches)
    section.header_distance = Inches(0)
    section.footer_distance = Inches(0)

    header = section.header
    header.is_linked_to_previous = False

    for paragraph in list(header.paragraphs):
        paragraph._p.getparent().remove(paragraph._p)

    page_w = section.page_width
    page_w_twips = int(page_w.emu // 635)
    margin_twips = int(section.left_margin.emu // 635)
    bleed_twips = margin_twips
    tbl_total_twips = page_w_twips + (bleed_twips * 2)

    table = header.add_table(rows=1, cols=1, width=page_w)
    table.autofit = False
    table.allow_autofit = False

    tbl_pr = table._tbl.tblPr
    for tag in ("w:tblW", "w:tblInd", "w:tblLayout"):
        for old in tbl_pr.findall(qn(tag)):
            tbl_pr.remove(old)

    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), str(tbl_total_twips))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_w)

    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), str(-(margin_twips + bleed_twips)))
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)

    tbl_layout = OxmlElement("w:tblLayout")
    tbl_layout.set(qn("w:type"), "fixed")
    tbl_pr.append(tbl_layout)

    cell = table.cell(0, 0)
    shade_cell(cell, "0D1B2A")
    remove_cell_borders(cell)
    set_cell_margins(cell, top=250, bottom=230, left=200, right=200)

    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(tbl_total_twips))
    tc_w.set(qn("w:type"), "dxa")

    name_p = cell.paragraphs[0]
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(name_p, before=0, after=0, line=1.0)
    shade_paragraph(name_p, "0D1B2A")
    add_header_background_shape(name_p, "0D1B2A")
    name_run = name_p.add_run(NAME)
    set_run(name_run, font=NAME_FONT, size=26, bold=True, color=WHITE)

    rule_p = cell.add_paragraph()
    rule_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(rule_p, before=3, after=3, line=1.0)
    shade_paragraph(rule_p, "0D1B2A")
    rule_run = rule_p.add_run("─" * 62)
    set_run(rule_run, font=NAME_FONT, size=8, color=GOLD)

    contact_p = cell.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(contact_p, before=2, after=0, line=1.15)
    shade_paragraph(contact_p, "0D1B2A")
    separator = "   |   "
    for index, (text, url) in enumerate(CONTACT_PARTS):
        if not text:
            continue
        if index > 0:
            sep_run = contact_p.add_run(separator)
            set_run(sep_run, font=CONTACT_FONT, size=9.5, color=GOLD)
        if url:
            add_hyperlink(contact_p, text, url, color=GOLD, font=CONTACT_FONT, size=9.5)
        else:
            run = contact_p.add_run(text)
            set_run(run, font=CONTACT_FONT, size=9.5, color=GOLD)

    footer = section.footer
    footer.is_linked_to_previous = False
    for paragraph in list(footer.paragraphs):
        paragraph._p.getparent().remove(paragraph._p)


def add_section_heading(doc, text):
    p = doc.add_paragraph()
    set_paragraph_format(p, before=8, after=2, line=1.1)
    run = p.add_run(text.upper())
    set_run(run, size=11.5, bold=True, color=STEEL)
    add_paragraph_bottom_border(p, color_hex="C9A84C", size=6)
    return p


def add_bullet(doc, text, *, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    set_paragraph_format(p, before=0, after=2, line=1.15)
    pf = p.paragraph_format
    pf.left_indent = Inches(0.18)
    pf.first_line_indent = Inches(-0.18)
    for run in p.runs:
        run.text = ""
    run = p.add_run(text)
    set_run(run, size=size, color=BLACK)
    return p


def add_job_block(doc, title, employer_line, dates):
    p1 = doc.add_paragraph()
    set_paragraph_format(p1, before=6, after=0, line=1.1)
    run = p1.add_run(title)
    set_run(run, size=11, bold=True, color=GOLD)

    p2 = doc.add_paragraph()
    set_paragraph_format(p2, before=0, after=2, line=1.1)
    employer_run = p2.add_run(employer_line)
    set_run(employer_run, size=10.5, italic=True, color=BLACK)
    date_run = p2.add_run("    " + dates)
    set_run(date_run, size=10, color=GRAY)
    return p1, p2


def new_document():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = Pt(10.5)
    return doc


def build_footer(document, *, show_page_numbers=False):
    """No-op by design. Troy's current locked standard uses no footer."""
    return None


def _append_page_x_of_y(paragraph):
    """No-op retained for backward compatibility with older scripts."""
    return None


__all__ = [
    "NAVY",
    "GOLD",
    "STEEL",
    "WHITE",
    "BLACK",
    "GRAY",
    "BODY_FONT",
    "NAME_FONT",
    "CONTACT_FONT",
    "NAME",
    "CONTACT_PARTS",
    "build_navy_header",
    "add_section_heading",
    "add_bullet",
    "add_job_block",
    "set_run",
    "set_paragraph_format",
    "add_hyperlink",
    "shade_cell",
    "set_cell_margins",
    "remove_cell_borders",
    "add_paragraph_bottom_border",
    "new_document",
    "build_footer",
]
