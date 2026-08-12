from docx import Document
from docx.oxml.ns import qn

from build_doc import EDUCATION_BLOCKS, build_document
from docx_header import build_navy_header, new_document


def test_resume_headings_and_subheadings_use_garamond(tmp_path):
    output = tmp_path / "heading-font-check.docx"
    build_document(
        "resume",
        {
            "summary": "Grounded summary.",
            "skills": "Grounded skills.",
            "jobs": [
                {
                    "title": "Investigator",
                    "employer": "Agency",
                    "dates": "2020 - 2024",
                    "bullets": ["Investigated documented matters."],
                }
            ],
            "extra_sections": [],
        },
        output,
    )

    document = Document(output)
    required = {
        "PROFESSIONAL SUMMARY",
        "CORE SKILLS",
        "EXPERIENCE",
        "Investigator",
        "EDUCATION",
        "Master of Arts, Police Leadership, Administration and Education",
        "Bachelor of Arts, Criminal Justice, Magna Cum Laude",
        "Associate of Arts, Criminal Justice, Magna Cum Laude",
    }
    observed = {}
    for paragraph in document.paragraphs:
        if paragraph.text in required:
            observed[paragraph.text] = {
                run.font.name for run in paragraph.runs if run.text.strip()
            }

    assert set(observed) == required
    assert all(fonts == {"Garamond"} for fonts in observed.values())

    all_runs = [run for paragraph in document.paragraphs for run in paragraph.runs if run.text.strip()]
    assert all(run.font.name == "Garamond" for run in all_runs)
    assert all(run.font.size.pt == 12 for run in all_runs)


def test_education_blocks_preserve_each_locked_gpa_and_year():
    assert EDUCATION_BLOCKS == (
        (
            "Master of Arts, Police Leadership, Administration and Education",
            "University of St. Thomas, St. Paul, MN",
            "GPA: 3.94",
            "2005",
        ),
        (
            "Bachelor of Arts, Criminal Justice, Magna Cum Laude",
            "St. Cloud State University, St. Cloud, MN",
            "GPA: 3.51",
            "1998",
        ),
        (
            "Associate of Arts, Criminal Justice, Magna Cum Laude",
            "St. Cloud State University, St. Cloud, MN",
            "GPA: 3.50",
            "1996",
        ),
    )


def test_header_paragraphs_have_navy_renderer_fallback():
    document = new_document()
    build_navy_header(document)
    section = document.sections[0]

    assert section.different_first_page_header_footer
    assert document.settings.odd_and_even_pages_header_footer
    for header in (section.header, section.first_page_header, section.even_page_header):
        assert len(header.tables) == 1
        cell = header.tables[0].cell(0, 0)
        shading = cell._tc.tcPr.find(qn("w:shd"))
        assert shading is not None
        assert shading.get(qn("w:fill")) == "0D1B2A"
        assert len(cell.paragraphs) == 3
