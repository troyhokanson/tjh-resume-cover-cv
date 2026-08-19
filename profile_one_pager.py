"""Shared one-page/profile layout built on the locked Troy document helpers."""

from __future__ import annotations

import json
from pathlib import Path

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from docx_header import (
    BLACK,
    NAME_FONT,
    STEEL,
    add_bullet,
    add_job_block,
    add_section_heading,
    build_navy_header,
    new_document,
    set_paragraph_format,
    set_run,
)


ROOT = Path(__file__).resolve().parent
CONTRACT_PATH = ROOT / "skills" / "build-troy-application" / "workflow_contract.json"


def _layout() -> dict[str, int | str]:
    contract = json.loads(CONTRACT_PATH.read_text(encoding="utf-8"))
    return contract["layout_minimums"]


LAYOUT = _layout()


def new_profile_document():
    """Create a branded one-pager that inherits the locked header implementation."""
    doc = new_document()
    build_navy_header(
        doc,
        body_top_margin_inches=1.55,
        body_bottom_margin_inches=0.55,
        body_left_margin_inches=0.6,
        body_right_margin_inches=0.6,
    )
    return doc


def add_profile_identity(doc, title: str, focus_lines: list[str]):
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(title_p, before=0, after=4, line=1.0)
    set_run(title_p.add_run(title.upper()), font=NAME_FONT, size=12, bold=True, color=STEEL)

    for index, focus in enumerate(focus_lines):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_format(p, before=0, after=8 if index == len(focus_lines) - 1 else 2, line=1.05)
        set_run(p.add_run(focus), size=9.5, bold=True, color=BLACK)
    return title_p


def add_profile_section(doc, text: str):
    p = add_section_heading(doc, text)
    for run in p.runs:
        run.font.name = NAME_FONT
    p.paragraph_format.space_before = Pt(int(LAYOUT["major_section_before"]) / 20)
    p.paragraph_format.space_after = Pt(int(LAYOUT["major_section_after"]) / 20)
    p.paragraph_format.keep_with_next = True
    return p


def add_profile_paragraph(doc, text: str, *, size: float = 9.5, after: float | None = None):
    p = doc.add_paragraph()
    after_points = int(LAYOUT["body_paragraph_after"]) / 20 if after is None else after
    set_paragraph_format(p, before=0, after=after_points, line=1.08)
    set_run(p.add_run(text), size=size, color=BLACK)
    return p


def add_profile_role(doc, title: str, employer: str, dates: str):
    title_p, employer_p = add_job_block(doc, title, employer, dates)
    title_p.paragraph_format.space_before = Pt(int(LAYOUT["major_entry_before"]) / 20)
    title_p.paragraph_format.space_after = Pt(int(LAYOUT["job_title_after"]) / 20)
    employer_p.paragraph_format.space_after = Pt(int(LAYOUT["employer_line_after"]) / 20)
    title_p.paragraph_format.keep_with_next = True
    employer_p.paragraph_format.keep_with_next = True
    return title_p, employer_p


def add_profile_bullet(doc, text: str, *, size: float = 9.0):
    p = add_bullet(doc, text, size=size)
    p.paragraph_format.space_after = Pt(int(LAYOUT["bullet_after"]) / 20)
    return p


__all__ = [
    "CONTRACT_PATH",
    "LAYOUT",
    "new_profile_document",
    "add_profile_identity",
    "add_profile_section",
    "add_profile_paragraph",
    "add_profile_role",
    "add_profile_bullet",
]
